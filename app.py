import json
import os
import re
import shutil
import sqlite3
import subprocess
import time
import tomllib
import uuid
import wave
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Literal, get_args

from docx import Document
from fastapi import BackgroundTasks, Body, FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from pypdf import PdfReader

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None


APP_DIR = Path(__file__).resolve().parent
STORAGE_DIR = APP_DIR / "storage"
UPLOADS_DIR = STORAGE_DIR / "uploads"
EXTRACTED_DIR = STORAGE_DIR / "extracted"
ANNOTATED_DIR = STORAGE_DIR / "annotated"
AUDIO_DIR = STORAGE_DIR / "audio"
DB_PATH = STORAGE_DIR / "app.db"
INDEX_PATH = APP_DIR / "index.html"
TTS_TAG_CONFIG_PATH = APP_DIR / "tts_tags.toml"

ALLOWED_SOURCE_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_ANNOTATION_EXTENSIONS = {".docx"}
CHAPTER_PATTERN = re.compile(
    r"^(chapter|part|book)\s+([0-9ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)\b",
    re.IGNORECASE,
)
INLINE_TTS_TAG_PATTERN = re.compile(r"(\[[^\]]+\])")
DEFAULT_VOICE_NAME = "Schedar"
DEFAULT_LANGUAGE_CODE = "en-US"
DEFAULT_AUDIO_SPEED = 1.05
DEFAULT_TTS_MAX_ATTEMPTS = 3
DEFAULT_AUDIO_SAMPLE_RATE = 24000
DEFAULT_STYLE_INSTRUCTION = "Read aloud in a warm, welcoming tone with clear diction and natural pacing."
DEFAULT_GEMINI_TEXT_MODEL = "gemini-2.5-flash"
DEFAULT_GEMINI_TTS_MODEL = "gemini-3.1-flash-tts-preview"
DEFAULT_ANNOTATION_BATCH_MAX_SEGMENTS = 45
DEFAULT_ANNOTATION_BATCH_MAX_CHARS = 12000
CHUNK_TARGET_PRESETS = [100, 200, 300, 400, 500, 650, 800, 1000, 1250, 1500]
CHUNK_HARD_OVERFLOW_PRESETS = [0, 25, 50, 75, 100, 150, 200, 300, 400, 500]
DEFAULT_MAX_CHARS = 500
DEFAULT_HARD_MAX_CHARS = 650
MAX_HARD_MAX_CHARS = CHUNK_TARGET_PRESETS[-1] + CHUNK_HARD_OVERFLOW_PRESETS[-1]
SENTENCE_BOUNDARY_PATTERN = re.compile(r"(?<=[.!?…])\s+")
CLAUSE_BOUNDARY_PATTERN = re.compile(r"(?<=[,;:])\s+")

SUPPORTED_GEMINI_VOICES = [
    {"name": "Zephyr", "description": "Bright", "gender": "Female"},
    {"name": "Puck", "description": "Upbeat", "gender": "Male"},
    {"name": "Charon", "description": "Informative", "gender": "Male"},
    {"name": "Kore", "description": "Firm", "gender": "Female"},
    {"name": "Fenrir", "description": "Excitable", "gender": "Male"},
    {"name": "Leda", "description": "Youthful", "gender": "Female"},
    {"name": "Orus", "description": "Firm", "gender": "Male"},
    {"name": "Aoede", "description": "Breezy", "gender": "Female"},
    {"name": "Callirrhoe", "description": "Easy-going", "gender": "Female"},
    {"name": "Autonoe", "description": "Bright", "gender": "Female"},
    {"name": "Enceladus", "description": "Breathy", "gender": "Male"},
    {"name": "Iapetus", "description": "Clear", "gender": "Male"},
    {"name": "Umbriel", "description": "Easy-going", "gender": "Male"},
    {"name": "Algieba", "description": "Smooth", "gender": "Male"},
    {"name": "Despina", "description": "Smooth", "gender": "Female"},
    {"name": "Erinome", "description": "Clear", "gender": "Female"},
    {"name": "Algenib", "description": "Gravelly", "gender": "Male"},
    {"name": "Rasalgethi", "description": "Informative", "gender": "Male"},
    {"name": "Laomedeia", "description": "Upbeat", "gender": "Female"},
    {"name": "Achernar", "description": "Soft", "gender": "Female"},
    {"name": "Alnilam", "description": "Firm", "gender": "Male"},
    {"name": "Schedar", "description": "Even", "gender": "Male"},
    {"name": "Gacrux", "description": "Mature", "gender": "Female"},
    {"name": "Pulcherrima", "description": "Forward", "gender": "Female"},
    {"name": "Achird", "description": "Friendly", "gender": "Male"},
    {"name": "Zubenelgenubi", "description": "Casual", "gender": "Male"},
    {"name": "Vindemiatrix", "description": "Gentle", "gender": "Female"},
    {"name": "Sadachbia", "description": "Lively", "gender": "Male"},
    {"name": "Sadaltager", "description": "Knowledgeable", "gender": "Male"},
    {"name": "Sulafat", "description": "Warm", "gender": "Female"},
]

SUPPORTED_LANGUAGE_OPTIONS = [
    {"code": "ar-XA", "label": "Arabic"},
    {"code": "bn-IN", "label": "Bangla"},
    {"code": "cmn-CN", "label": "Chinese (Mandarin)"},
    {"code": "nl-NL", "label": "Dutch"},
    {"code": "en-AU", "label": "English (Australia)"},
    {"code": "en-GB", "label": "English (UK)"},
    {"code": "en-IN", "label": "English (India)"},
    {"code": "en-US", "label": "English (US)"},
    {"code": "fr-CA", "label": "French (Canada)"},
    {"code": "fr-FR", "label": "French (France)"},
    {"code": "de-DE", "label": "German"},
    {"code": "gu-IN", "label": "Gujarati"},
    {"code": "hi-IN", "label": "Hindi"},
    {"code": "id-ID", "label": "Indonesian"},
    {"code": "it-IT", "label": "Italian"},
    {"code": "ja-JP", "label": "Japanese"},
    {"code": "kn-IN", "label": "Kannada"},
    {"code": "ko-KR", "label": "Korean"},
    {"code": "ml-IN", "label": "Malayalam"},
    {"code": "mr-IN", "label": "Marathi"},
    {"code": "pl-PL", "label": "Polish"},
    {"code": "pt-BR", "label": "Portuguese (Brazil)"},
    {"code": "ru-RU", "label": "Russian"},
    {"code": "es-ES", "label": "Spanish (Spain)"},
    {"code": "es-US", "label": "Spanish (US)"},
    {"code": "ta-IN", "label": "Tamil"},
    {"code": "te-IN", "label": "Telugu"},
    {"code": "th-TH", "label": "Thai"},
    {"code": "tr-TR", "label": "Turkish"},
    {"code": "vi-VN", "label": "Vietnamese"},
]

if load_dotenv is not None:
    load_dotenv(APP_DIR / ".env")

SegmentType = Literal["heading", "narration", "dialogue", "quote", "front_matter", "back_matter", "other"]
StyleTag = Literal[
    "neutral",
    "serious",
    "warm",
    "curious",
    "tense",
    "sad",
    "excited",
    "soft",
    "whisper",
    "shout",
    "tired",
    "amazed",
    "trembling",
    "panicked",
    "sarcastic",
    "reluctant",
]
PaceTag = Literal["slow", "steady", "fast"]
PauseTag = Literal["none", "short", "medium", "long"]
TagGenerationMode = Literal["off", "conservative", "balanced", "expressive"]
AudioOutputFormat = Literal["mp3", "wav"]
AudioQualityPreset = Literal["best", "standard"]
PreChapterTextMode = Literal["attach_to_chapter_1", "separate_chapter_0"]

DEFAULT_EXPRESSIVE_TAG_MODE: TagGenerationMode = "conservative"
DEFAULT_VOCALIZATION_TAG_MODE: TagGenerationMode = "off"
TAG_GENERATION_MODE_RANKS = {
    "off": 0,
    "conservative": 1,
    "balanced": 2,
    "expressive": 3,
}

DEFAULT_TTS_TAG_CONFIG = {
    "expressive_tags": [
        {"tag": "[frustrated]", "min_mode": "conservative"},
        {"tag": "[annoyed]", "min_mode": "conservative"},
        {"tag": "[angry]", "min_mode": "expressive"},
        {"tag": "[tense]", "min_mode": "conservative"},
        {"tag": "[agitated]", "min_mode": "expressive"},
        {"tag": "[warm]", "min_mode": "conservative"},
        {"tag": "[weary]", "min_mode": "conservative"},
        {"tag": "[hesitant]", "min_mode": "conservative"},
        {"tag": "[relieved]", "min_mode": "conservative"},
        {"tag": "[reassuring]", "min_mode": "conservative"},
        {"tag": "[pleading]", "min_mode": "conservative"},
        {"tag": "[breathless]", "min_mode": "expressive"},
        {"tag": "[shaky]", "min_mode": "expressive"},
        {"tag": "[softly]", "min_mode": "conservative"},
        {"tag": "[afraid]", "min_mode": "conservative"},
        {"tag": "[disgusted]", "min_mode": "expressive"},
        {"tag": "[under her breath]", "min_mode": "conservative"},
        {"tag": "[through gritted teeth]", "min_mode": "expressive"},
        {"tag": "[voice breaking]", "min_mode": "conservative"},
        {"tag": "[after a pause]", "min_mode": "conservative"},
    ],
    "vocalization_tags": [
        {"tag": "[sighs]", "min_mode": "conservative"},
        {"tag": "[gasps]", "min_mode": "conservative"},
        {"tag": "[laughs]", "min_mode": "conservative"},
        {"tag": "[chuckles]", "min_mode": "conservative"},
        {"tag": "[snorts]", "min_mode": "expressive"},
        {"tag": "[scoffs]", "min_mode": "expressive"},
        {"tag": "[harumphs]", "min_mode": "expressive"},
        {"tag": "[huffs]", "min_mode": "conservative"},
        {"tag": "[breathing hard]", "min_mode": "conservative"},
        {"tag": "[exhales slowly]", "min_mode": "conservative"},
        {"tag": "[groans]", "min_mode": "expressive"},
        {"tag": "[moans]", "min_mode": "expressive"},
        {"tag": "[yelps]", "min_mode": "conservative"},
        {"tag": "[murmurs]", "min_mode": "conservative"},
        {"tag": "[whispers]", "min_mode": "conservative"},
        {"tag": "[coughs]", "min_mode": "conservative"},
        {"tag": "[sniffling]", "min_mode": "conservative"},
        {"tag": "[wincing]", "min_mode": "conservative"},
        {"tag": "[grunts]", "min_mode": "conservative"},
        {"tag": "[humming]", "min_mode": "conservative"},
        {"tag": "[drawls]", "min_mode": "conservative"},
        {"tag": "[shuddering]", "min_mode": "conservative"},
        {"tag": "[trembling breath]", "min_mode": "conservative"},
        {"tag": "[half laughing]", "min_mode": "conservative"},
        {"tag": "[half crying]", "min_mode": "conservative"},
        {"tag": "[gasping]", "min_mode": "conservative"},
        {"tag": "[swearing under breath]", "min_mode": "expressive"},
        {"tag": "[laughing nervously]", "min_mode": "conservative"},
        {"tag": "[fighting tears]", "min_mode": "expressive"},
        {"tag": "[breathing unevenly]", "min_mode": "conservative"},
    ],
}


def clone_default_tts_tag_config() -> dict[str, list[dict[str, str]]]:
    return {
        category: [dict(entry) for entry in entries]
        for category, entries in DEFAULT_TTS_TAG_CONFIG.items()
    }


def normalize_tts_tag_entries(raw_entries: Any, category_name: str) -> list[dict[str, str]]:
    if not isinstance(raw_entries, list):
        raise ValueError(f"{category_name} must be a list of tables")

    normalized: list[dict[str, str]] = []
    seen_tags: set[str] = set()
    for index, entry in enumerate(raw_entries, start=1):
        if not isinstance(entry, dict):
            raise ValueError(f"{category_name}[{index}] must be a table")
        tag = entry.get("tag")
        min_mode = entry.get("min_mode", DEFAULT_EXPRESSIVE_TAG_MODE)
        if not isinstance(tag, str) or not tag.startswith("[") or not tag.endswith("]"):
            raise ValueError(f"{category_name}[{index}].tag must be a bracketed string")
        if min_mode not in {"conservative", "balanced", "expressive"}:
            raise ValueError(
                f"{category_name}[{index}].min_mode must be conservative, balanced, or expressive"
            )
        if tag in seen_tags:
            raise ValueError(f"Duplicate tag in {category_name}: {tag}")
        normalized.append({"tag": tag.strip(), "min_mode": min_mode})
        seen_tags.add(tag)
    return normalized


def load_tts_tag_config() -> dict[str, list[dict[str, str]]]:
    config = clone_default_tts_tag_config()
    if not TTS_TAG_CONFIG_PATH.exists():
        return config

    try:
        with TTS_TAG_CONFIG_PATH.open("rb") as handle:
            parsed = tomllib.load(handle)
        config["expressive_tags"] = normalize_tts_tag_entries(parsed.get("expressive_tags", []), "expressive_tags")
        config["vocalization_tags"] = normalize_tts_tag_entries(parsed.get("vocalization_tags", []), "vocalization_tags")
    except Exception as exc:
        print(f"Failed to load {TTS_TAG_CONFIG_PATH.name}; using built-in tag defaults. {exc}")
        return clone_default_tts_tag_config()

    return config


TTS_TAG_CONFIG = load_tts_tag_config()

EXPRESSIVE_INLINE_TAGS = [entry["tag"] for entry in TTS_TAG_CONFIG["expressive_tags"]]
VOCALIZATION_INLINE_TAGS = [entry["tag"] for entry in TTS_TAG_CONFIG["vocalization_tags"]]

EXPRESSIVE_INLINE_TAG_SET = set(EXPRESSIVE_INLINE_TAGS)
VOCALIZATION_INLINE_TAG_SET = set(VOCALIZATION_INLINE_TAGS)
EXPRESSIVE_TAG_MIN_MODE = {
    entry["tag"]: entry["min_mode"] for entry in TTS_TAG_CONFIG["expressive_tags"]
}
VOCALIZATION_TAG_MIN_MODE = {
    entry["tag"]: entry["min_mode"] for entry in TTS_TAG_CONFIG["vocalization_tags"]
}

ALLOWED_GEMINI_TTS_TAGS = {
    "[short pause]",
    "[very fast]",
    "[very slow]",
    "[excitedly]",
    "[bored]",
    "[reluctantly]",
    "[sarcastically, one painfully slow word at a time]",
    "[sarcastically]",
    "[whisper]",
    "[whispers]",
    "[shouting]",
    "[sarcastic]",
    "[serious]",
    "[tired]",
    "[amazed]",
    "[crying]",
    "[curious]",
    "[excited]",
    "[mischievously]",
    "[panicked]",
    "[trembling]",
    "[yawn]",
    "[cough]",
    "[sighs]",
    "[gasp]",
    "[giggles]",
    "[laughs]",
}

STYLE_TO_GEMINI_TAG = {
    "neutral": None,
    "serious": "[serious]",
    "warm": None,
    "curious": "[curious]",
    "tense": None,
    "sad": None,
    "excited": "[excitedly]",
    "soft": None,
    "whisper": "[whispers]",
    "shout": "[shouting]",
    "tired": "[tired]",
    "amazed": "[amazed]",
    "trembling": "[trembling]",
    "panicked": "[panicked]",
    "sarcastic": "[sarcastically]",
    "reluctant": "[reluctantly]",
}

PAUSE_TO_GEMINI_TAG = {
    "none": None,
    "short": "[short pause]",
    "medium": None,
    "long": None,
}

PACE_TO_GEMINI_TAG = {
    "slow": "[very slow]",
    "steady": None,
    "fast": "[very fast]",
}

SHOUT_CUES = (
    "shouted",
    "yelled",
    "screamed",
    "roared",
    "bellowed",
    "shrieked",
    "cried out",
    "called out",
)

WHISPER_CUES = (
    "whispered",
    "murmured",
    "muttered",
    "hushed",
    "under his breath",
    "under her breath",
    "under their breath",
)

SERIOUS_CUES = (
    "seriously",
    "sternly",
    "grimly",
    "gravely",
    "solemnly",
    "quietly said",
)

TIRED_CUES = (
    "tired",
    "weary",
    "exhausted",
    "drained",
    "sleepy",
    "yawned",
    "worn out",
)

AMAZED_CUES = (
    "amazed",
    "astonished",
    "astounded",
    "incredible",
    "unbelievable",
    "wow",
    "oh my god",
    "oh my goodness",
)

TREMBLING_CUES = (
    "trembling",
    "trembled",
    "shaking",
    "quivering",
    "shivering",
    "voice shook",
    "hands shook",
)

PANICKED_CUES = (
    "panicked",
    "frantic",
    "in a panic",
    "hurry",
    "run",
    "help",
    "oh no",
)

SARCASTIC_CUES = (
    "sarcastically",
    "sarcastic",
    "dryly",
    "mockingly",
    "yeah, right",
    "as if",
)

RELUCTANT_CUES = (
    "reluctantly",
    "hesitant",
    "hesitantly",
    "unwilling",
    "finally admitted",
    "after a pause",
)

CURIOUS_QUESTION_CUES = (
    "what is",
    "what's",
    "where is",
    "where's",
    "who is",
    "who's",
    "how does",
    "how do",
    "why does",
    "did you see",
    "do you know",
    "really?",
    "is that",
    "asked",
    "ask",
    "wondered",
    "wonder",
    "wanted to know",
    "curiously",
)

SERIOUS_QUESTION_CUES = (
    "what happened",
    "what have you done",
    "is this true",
    "what are we going to do",
)

FAST_CUES = ("quickly", "hurried", "rushed", "rapidly")
SLOW_CUES = ("slowly", "hesitated", "drawled", "lingered")
OPENING_QUOTE_CHARS = ('"', '“')


def extract_inline_tts_tags(text: str) -> list[str]:
    return INLINE_TTS_TAG_PATTERN.findall(text)


def validate_inline_tts_tags(text: str) -> None:
    return None


def strip_inline_tts_tags(text: str) -> str:
    return INLINE_TTS_TAG_PATTERN.sub("", text).strip()


def find_inline_speech_start(text: str) -> int | None:
    positions = [text.find(character) for character in OPENING_QUOTE_CHARS if text.find(character) != -1]
    if not positions:
        return None
    return min(positions)


def segment_has_inline_speech(segment: "SegmentAnnotation") -> bool:
    return find_inline_speech_start(segment.clean_text) is not None


def segment_supports_inline_tags(segment: "SegmentAnnotation") -> bool:
    return segment.segment_type in {"dialogue", "quote"} or segment_has_inline_speech(segment)


def tag_generation_mode_rank(mode: TagGenerationMode) -> int:
    return TAG_GENERATION_MODE_RANKS[mode]


def allowed_inline_tags_for_mode(mode: TagGenerationMode, expressive_category: bool) -> list[str]:
    if mode == "off":
        return []

    tag_min_modes = EXPRESSIVE_TAG_MIN_MODE if expressive_category else VOCALIZATION_TAG_MIN_MODE
    return [
        tag
        for tag, min_mode in tag_min_modes.items()
        if tag_generation_mode_rank(mode) >= TAG_GENERATION_MODE_RANKS[min_mode]
    ]


def map_legacy_immersiveness_to_mode(level: str | None) -> TagGenerationMode:
    return {
        "low": "conservative",
        "medium": "balanced",
        "high": "expressive",
    }.get((level or "").strip().lower(), DEFAULT_EXPRESSIVE_TAG_MODE)


def normalize_generated_annotation_settings(extracted: dict[str, Any]) -> dict[str, TagGenerationMode]:
    settings = extracted.get("generated_annotation_settings")
    if isinstance(settings, dict):
        expressive_mode = settings.get("expressive_mode", DEFAULT_EXPRESSIVE_TAG_MODE)
        vocalization_mode = settings.get("vocalization_mode", DEFAULT_VOCALIZATION_TAG_MODE)
        return {
            "expressive_mode": expressive_mode if expressive_mode in get_args(TagGenerationMode) else DEFAULT_EXPRESSIVE_TAG_MODE,
            "vocalization_mode": vocalization_mode if vocalization_mode in get_args(TagGenerationMode) else DEFAULT_VOCALIZATION_TAG_MODE,
        }

    legacy_mode = map_legacy_immersiveness_to_mode(extracted.get("generated_annotation_immersiveness"))
    return {
        "expressive_mode": legacy_mode,
        "vocalization_mode": DEFAULT_VOCALIZATION_TAG_MODE,
    }


def segment_has_explicit_text(segment: "SegmentAnnotation", cues: tuple[str, ...]) -> bool:
    return text_contains_any(segment.clean_text.lower(), cues)


def tag_allowed_for_mode(tag: str, mode: TagGenerationMode, expressive_category: bool) -> bool:
    if mode == "off":
        return False
    min_mode = (EXPRESSIVE_TAG_MIN_MODE if expressive_category else VOCALIZATION_TAG_MIN_MODE).get(tag)
    if min_mode is None:
        return False
    return tag_generation_mode_rank(mode) >= TAG_GENERATION_MODE_RANKS[min_mode]


def infer_generated_expressive_tag(segment: "SegmentAnnotation", mode: TagGenerationMode) -> str | None:
    if mode == "off" or not segment_supports_inline_tags(segment):
        return None

    text = segment.clean_text
    lowered = text.lower()
    rank = tag_generation_mode_rank(mode)

    direct_candidates: list[tuple[str, bool]] = [
        ("[under her breath]", text_contains_any(lowered, WHISPER_CUES)),
        ("[through gritted teeth]", "gritted teeth" in lowered),
        ("[after a pause]", "after a pause" in lowered or "paused" in lowered),
        ("[voice breaking]", "voice breaking" in lowered or "voice broke" in lowered or "voice cracked" in lowered),
        ("[warm]", "warmly" in lowered or "kindly" in lowered or "gently" in lowered),
        ("[weary]", text_contains_any(lowered, TIRED_CUES)),
        ("[hesitant]", text_contains_any(lowered, RELUCTANT_CUES)),
        ("[relieved]", "relieved" in lowered or "thank goodness" in lowered or "thank god" in lowered),
        ("[reassuring]", "reassured" in lowered or "it's okay" in lowered or "you are safe" in lowered),
        ("[pleading]", "pleaded" in lowered or "begged" in lowered or "implored" in lowered),
        ("[breathless]", "breathless" in lowered or "panting" in lowered),
        ("[shaky]", text_contains_any(lowered, TREMBLING_CUES)),
        ("[softly]", "softly" in lowered or "gently" in lowered),
        ("[afraid]", "afraid" in lowered or "scared" in lowered or "terrified" in lowered or "fear" in lowered),
        ("[disgusted]", "disgusted" in lowered or "revolted" in lowered or "gross" in lowered),
        ("[frustrated]", "frustrated" in lowered or "exasperated" in lowered),
        ("[annoyed]", "annoyed" in lowered or "irritated" in lowered),
        ("[angry]", "angrily" in lowered or "furious" in lowered or "rage" in lowered),
        ("[tense]", "tense" in lowered or "stiffly" in lowered or "taut" in lowered),
        ("[agitated]", "agitated" in lowered or "restless" in lowered),
    ]
    for tag, matched in direct_candidates:
        if matched and tag_allowed_for_mode(tag, mode, expressive_category=True):
            return tag

    if rank >= 2:
        if infer_supported_style(segment) == "excited" and tag_allowed_for_mode("[agitated]", mode, expressive_category=True):
            return "[agitated]"
        if infer_supported_style(segment) == "reluctant" and tag_allowed_for_mode("[hesitant]", mode, expressive_category=True):
            return "[hesitant]"
        if infer_supported_style(segment) == "tired" and tag_allowed_for_mode("[weary]", mode, expressive_category=True):
            return "[weary]"
        if infer_supported_style(segment) == "trembling" and tag_allowed_for_mode("[shaky]", mode, expressive_category=True):
            return "[shaky]"
        if infer_supported_style(segment) == "panicked" and tag_allowed_for_mode("[afraid]", mode, expressive_category=True):
            return "[afraid]"
        if infer_supported_style(segment) == "serious" and tag_allowed_for_mode("[tense]", mode, expressive_category=True):
            return "[tense]"
        if infer_supported_style(segment) == "amazed" and tag_allowed_for_mode("[breathless]", mode, expressive_category=True):
            return "[breathless]"
        if infer_supported_style(segment) == "sarcastic" and tag_allowed_for_mode("[annoyed]", mode, expressive_category=True):
            return "[annoyed]"

    if rank >= 3:
        if looks_like_shouting(text) and tag_allowed_for_mode("[angry]", mode, expressive_category=True):
            return "[angry]"
        if "!" in text and "?" in text and tag_allowed_for_mode("[agitated]", mode, expressive_category=True):
            return "[agitated]"
        if text.count("!") >= 2 and tag_allowed_for_mode("[frustrated]", mode, expressive_category=True):
            return "[frustrated]"
    return None


def infer_generated_vocalization_tag(segment: "SegmentAnnotation", mode: TagGenerationMode) -> str | None:
    if mode == "off" or not segment_supports_inline_tags(segment):
        return None

    text = segment.clean_text.lower()
    rank = tag_generation_mode_rank(mode)

    direct_candidates: list[tuple[str, bool]] = [
        ("[sighs]", "sighed" in text or "sighs" in text or "sigh" in text),
        ("[gasps]", "gasped" in text or "gasps" in text or "gasp" in text),
        ("[laughs]", "laughed" in text or "laughs" in text or "laughing" in text),
        ("[chuckles]", "chuckled" in text or "chuckles" in text or "chuckle" in text),
        ("[murmurs]", "murmured" in text or "murmurs" in text),
        ("[whispers]", text_contains_any(text, WHISPER_CUES)),
        ("[coughs]", "coughed" in text or "coughs" in text or "cough" in text),
        ("[sniffling]", "sniffling" in text or "sniffled" in text or "sniffles" in text),
        ("[trembling breath]", "trembling breath" in text or "shaky breath" in text),
        ("[breathing unevenly]", "breathing unevenly" in text or "ragged breath" in text),
        ("[wincing]", "wincing" in text or "winced" in text),
        ("[laughing nervously]", "laughed nervously" in text or "laughing nervously" in text),
        ("[fighting tears]", "fighting tears" in text or "choked up" in text),
    ]
    for tag, matched in direct_candidates:
        if matched and tag_allowed_for_mode(tag, mode, expressive_category=False):
            return tag

    if rank >= 2:
        if infer_supported_style(segment) == "amazed" and tag_allowed_for_mode("[gasps]", mode, expressive_category=False):
            return "[gasps]"
        if infer_supported_style(segment) == "whisper" and tag_allowed_for_mode("[whispers]", mode, expressive_category=False):
            return "[whispers]"
        if infer_supported_style(segment) == "trembling" and tag_allowed_for_mode("[trembling breath]", mode, expressive_category=False):
            return "[trembling breath]"
        if infer_supported_style(segment) == "panicked" and tag_allowed_for_mode("[breathing unevenly]", mode, expressive_category=False):
            return "[breathing unevenly]"
        if infer_supported_pace(segment) == "slow" and "..." in segment.clean_text and tag_allowed_for_mode("[drawls]", mode, expressive_category=False):
            return "[drawls]"

    if rank >= 3:
        if infer_supported_style(segment) == "sarcastic" and tag_allowed_for_mode("[scoffs]", mode, expressive_category=False):
            return "[scoffs]"
        if looks_like_shouting(segment.clean_text) and tag_allowed_for_mode("[huffs]", mode, expressive_category=False):
            return "[huffs]"
    return None


def sanitize_generated_inline_tags(
    source_segment: "SegmentAnnotation",
    inline_tags: list[str] | None,
    expressive_mode: TagGenerationMode,
    vocalization_mode: TagGenerationMode,
) -> list[str]:
    if not inline_tags or not segment_supports_inline_tags(source_segment):
        return []

    sanitized: list[str] = []
    seen: set[str] = set()
    expressive_used = False
    vocalization_used = False
    for raw_tag in inline_tags:
        tag = raw_tag.strip()
        if not tag or tag in seen:
            continue
        if tag in EXPRESSIVE_INLINE_TAG_SET:
            if expressive_used or not tag_allowed_for_mode(tag, expressive_mode, expressive_category=True):
                continue
            expressive_used = True
        elif tag in VOCALIZATION_INLINE_TAG_SET:
            if vocalization_used or not tag_allowed_for_mode(tag, vocalization_mode, expressive_category=False):
                continue
            vocalization_used = True
        else:
            continue
        sanitized.append(tag)
        seen.add(tag)
        break
    return sanitized


def resolve_segment_inline_tags(
    segment: "SegmentAnnotation",
    expressive_mode: TagGenerationMode,
    vocalization_mode: TagGenerationMode,
) -> list[str]:
    explicit_tags = sanitize_generated_inline_tags(
        segment,
        segment.inline_tags,
        expressive_mode,
        vocalization_mode,
    )
    if explicit_tags:
        return explicit_tags

    vocalization_tag = infer_generated_vocalization_tag(segment, vocalization_mode)
    expressive_tag = infer_generated_expressive_tag(segment, expressive_mode)
    if vocalization_tag:
        return [vocalization_tag]
    if expressive_tag:
        return [expressive_tag]
    return []


def text_contains_any(text: str, cues: tuple[str, ...]) -> bool:
    return any(cue in text for cue in cues)


def looks_like_shouting(text: str) -> bool:
    letters = [character for character in text if character.isalpha()]
    if letters:
        uppercase_ratio = sum(1 for character in letters if character.isupper()) / len(letters)
        if len(letters) >= 8 and uppercase_ratio >= 0.7:
            return True
    lowered = text.lower()
    return text.count("!") >= 2 or text_contains_any(lowered, SHOUT_CUES)


def infer_supported_style(segment: "SegmentAnnotation") -> StyleTag:
    if not segment_supports_inline_tags(segment):
        return "neutral"

    text = segment.clean_text.strip()
    lowered = text.lower()

    if looks_like_shouting(text):
        return "shout"
    if text_contains_any(lowered, WHISPER_CUES):
        return "whisper"
    if text_contains_any(lowered, PANICKED_CUES) and ("!" in text or "?" in text):
        return "panicked"
    if text_contains_any(lowered, TREMBLING_CUES):
        return "trembling"
    if text_contains_any(lowered, AMAZED_CUES):
        return "amazed"
    if text_contains_any(lowered, TIRED_CUES):
        return "tired"
    if text_contains_any(lowered, RELUCTANT_CUES):
        return "reluctant"
    if text_contains_any(lowered, SARCASTIC_CUES):
        return "sarcastic"
    if "?" in text and text_contains_any(lowered, CURIOUS_QUESTION_CUES):
        return "curious"
    if "?" in text and text_contains_any(lowered, SERIOUS_QUESTION_CUES):
        return "serious"
    if "?" in text and len(text) <= 90:
        return "curious"
    if "!" in text:
        return "excited"
    if text_contains_any(lowered, SERIOUS_CUES):
        return "serious"
    return "neutral"


def style_is_supported_by_text(segment: "SegmentAnnotation", style: StyleTag) -> bool:
    if style == "neutral":
        return True
    return infer_supported_style(segment) == style


def infer_supported_pace(segment: "SegmentAnnotation") -> PaceTag:
    if not segment_supports_inline_tags(segment):
        return "steady"

    lowered = segment.clean_text.lower()
    if text_contains_any(lowered, FAST_CUES) or segment.clean_text.count("!") >= 2:
        return "fast"
    if text_contains_any(lowered, SLOW_CUES) or "..." in segment.clean_text:
        return "slow"
    return "steady"


def pick_inline_style_tag(segment: "SegmentAnnotation", mode: TagGenerationMode) -> str | None:
    supported_style = infer_supported_style(segment)
    if mode == "off":
        return None

    allowed_styles = {
        "shout",
        "whisper",
        "curious",
        "excited",
        "serious",
        "tired",
        "amazed",
        "reluctant",
        "sarcastic",
        "trembling",
        "panicked",
    }

    if supported_style not in allowed_styles:
        return None
    return STYLE_TO_GEMINI_TAG.get(supported_style)


def pick_inline_pace_tag(segment: "SegmentAnnotation", mode: TagGenerationMode) -> str | None:
    if mode not in {"balanced", "expressive"}:
        return None
    pace = infer_supported_pace(segment)
    return PACE_TO_GEMINI_TAG.get(pace)


def apply_inline_tag_to_text(text: str, tag: str) -> str:
    speech_start = find_inline_speech_start(text)
    if speech_start is None or speech_start == 0:
        return f"{tag} {text}".strip()
    return f"{text[:speech_start]}{tag} {text[speech_start:]}".strip()


def build_inline_tts_text(
    segment: "SegmentAnnotation",
    expressive_mode: TagGenerationMode = DEFAULT_EXPRESSIVE_TAG_MODE,
    vocalization_mode: TagGenerationMode = DEFAULT_VOCALIZATION_TAG_MODE,
) -> str:
    if not segment.speak:
        return segment.clean_text.strip()
    if not segment_supports_inline_tags(segment):
        return segment.clean_text.strip()
    if extract_inline_tts_tags(segment.clean_text):
        return segment.clean_text.strip()

    tags = [tag for tag in resolve_segment_inline_tags(segment, expressive_mode, vocalization_mode) if tag in ALLOWED_GEMINI_TTS_TAGS or tag in EXPRESSIVE_INLINE_TAG_SET or tag in VOCALIZATION_INLINE_TAG_SET]
    style_tag = pick_inline_style_tag(segment, expressive_mode)
    pace_tag = pick_inline_pace_tag(segment, expressive_mode)

    if not tags:
        if style_tag and style_tag in ALLOWED_GEMINI_TTS_TAGS:
            tags.append(style_tag)
        elif pace_tag and pace_tag in ALLOWED_GEMINI_TTS_TAGS:
            tags.append(pace_tag)

    tagged_text = segment.clean_text.strip()
    for tag in tags:
        tagged_text = apply_inline_tag_to_text(tagged_text, tag)
    return tagged_text


class ChunkPreviewRequest(BaseModel):
    max_chars: int = Field(default=DEFAULT_MAX_CHARS, ge=100, le=CHUNK_TARGET_PRESETS[-1])
    hard_max_chars: int = Field(default=DEFAULT_HARD_MAX_CHARS, ge=100, le=MAX_HARD_MAX_CHARS)
    read_chapter_titles: bool = True
    pre_chapter_text_mode: PreChapterTextMode = "attach_to_chapter_1"


class DraftAnnotationRequest(BaseModel):
    expressive_mode: TagGenerationMode = DEFAULT_EXPRESSIVE_TAG_MODE
    vocalization_mode: TagGenerationMode = DEFAULT_VOCALIZATION_TAG_MODE


class TtsPromptPreviewRequest(BaseModel):
    max_chars: int = Field(default=DEFAULT_MAX_CHARS, ge=100, le=CHUNK_TARGET_PRESETS[-1])
    hard_max_chars: int = Field(default=DEFAULT_HARD_MAX_CHARS, ge=100, le=MAX_HARD_MAX_CHARS)
    read_chapter_titles: bool = True
    pre_chapter_text_mode: PreChapterTextMode = "attach_to_chapter_1"
    chapter_index: int = Field(ge=0)
    chunk_index: int = Field(ge=0)
    voice_name: str = Field(default=DEFAULT_VOICE_NAME, min_length=1, max_length=100)
    language_code: str = Field(default=DEFAULT_LANGUAGE_CODE, min_length=2, max_length=20)
    style_instruction: str = Field(default=DEFAULT_STYLE_INSTRUCTION, min_length=1, max_length=500)
    speech_rate: float = Field(default=DEFAULT_AUDIO_SPEED, ge=0.8, le=1.5)
    output_format: AudioOutputFormat = "mp3"
    audio_quality: AudioQualityPreset = "best"


class GenerateChunkAudioRequest(TtsPromptPreviewRequest):
    pass


class GenerateBookAudioRequest(BaseModel):
    max_chars: int = Field(default=DEFAULT_MAX_CHARS, ge=100, le=CHUNK_TARGET_PRESETS[-1])
    hard_max_chars: int = Field(default=DEFAULT_HARD_MAX_CHARS, ge=100, le=MAX_HARD_MAX_CHARS)
    read_chapter_titles: bool = True
    pre_chapter_text_mode: PreChapterTextMode = "attach_to_chapter_1"
    voice_name: str = Field(default=DEFAULT_VOICE_NAME, min_length=1, max_length=100)
    language_code: str = Field(default=DEFAULT_LANGUAGE_CODE, min_length=2, max_length=20)
    style_instruction: str = Field(default=DEFAULT_STYLE_INSTRUCTION, min_length=1, max_length=500)
    speech_rate: float = Field(default=DEFAULT_AUDIO_SPEED, ge=0.8, le=1.5)
    output_format: AudioOutputFormat = "mp3"
    audio_quality: AudioQualityPreset = "best"
    chapter_indexes: list[int] | None = None


class TranscriptCoverageReview(BaseModel):
    is_complete: bool
    summary: str
    missing_excerpts: list[str] = Field(default_factory=list)


class SegmentAnnotation(BaseModel):
    segment_id: str
    segment_type: SegmentType
    speak: bool
    clean_text: str
    join_with_previous: bool = False
    style: StyleTag
    pace: PaceTag
    pause_after: PauseTag
    inline_tags: list[str] = Field(default_factory=list)
    speaker_hint: str | None = None
    notes: str | None = None


class AnnotatedChapter(BaseModel):
    chapter_id: str
    chapter_title: str
    leading_segment_count: int = 0
    segments: list[SegmentAnnotation]


class AnnotationDocument(BaseModel):
    title: str
    voice_name: str = DEFAULT_VOICE_NAME
    chapters: list[AnnotatedChapter]


def build_default_generator_settings() -> dict[str, Any]:
    return {
        "voice_name": DEFAULT_VOICE_NAME,
        "language_code": DEFAULT_LANGUAGE_CODE,
        "style_instruction": DEFAULT_STYLE_INSTRUCTION,
        "speech_rate": DEFAULT_AUDIO_SPEED,
        "max_chars": DEFAULT_MAX_CHARS,
        "hard_max_chars": DEFAULT_HARD_MAX_CHARS,
        "read_chapter_titles": True,
        "pre_chapter_text_mode": "attach_to_chapter_1",
        "output_format": "mp3",
        "audio_quality": "best",
        "chapter_indexes": [],
    }


def get_generator_settings(extracted: dict[str, Any]) -> dict[str, Any]:
    settings = build_default_generator_settings()
    saved_settings = extracted.get("generator_settings") or {}
    settings.update(saved_settings)
    settings["max_chars"], settings["hard_max_chars"] = normalize_chunk_limits(
        settings.get("max_chars", DEFAULT_MAX_CHARS),
        settings.get("hard_max_chars"),
    )
    return settings


def get_supported_language_codes() -> set[str]:
    return {option["code"] for option in SUPPORTED_LANGUAGE_OPTIONS}


def validate_language_code(language_code: str) -> str:
    if language_code not in get_supported_language_codes():
        raise HTTPException(status_code=400, detail=f"Unsupported language code: {language_code}")
    return language_code


def build_chapter_summaries(annotations: AnnotationDocument) -> list[dict[str, Any]]:
    return build_chapter_summaries_from_chapters(annotations.chapters)


def build_chapter_summaries_from_chapters(chapters: list[AnnotatedChapter]) -> list[dict[str, Any]]:
    summaries: list[dict[str, Any]] = []
    for chapter_index, chapter in enumerate(chapters):
        speakable_segments = [segment for segment in chapter.segments if segment.speak]
        summaries.append(
            {
                "chapter_index": chapter_index,
                "chapter_title": chapter.chapter_title,
                "segment_count": len(chapter.segments),
                "speakable_segment_count": len(speakable_segments),
                "char_count": sum(len(segment.clean_text) for segment in speakable_segments),
            }
        )
    return summaries


def normalize_selected_chapter_indexes(chapter_count: int, chapter_indexes: list[int] | None) -> list[int]:
    if not chapter_indexes:
        return list(range(chapter_count))

    normalized = sorted(set(chapter_indexes))
    invalid_indexes = [index for index in normalized if index < 0 or index >= chapter_count]
    if invalid_indexes:
        raise HTTPException(status_code=400, detail=f"Invalid chapter indexes requested: {invalid_indexes}")
    return normalized


app = FastAPI(title="AutoAudioBook")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


def ensure_storage() -> None:
    for path in [STORAGE_DIR, UPLOADS_DIR, EXTRACTED_DIR, ANNOTATED_DIR, AUDIO_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def get_connection() -> sqlite3.Connection:
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    return connection


def init_db() -> None:
    ensure_storage()
    with get_connection() as connection:
        connection.executescript(
            """
            CREATE TABLE IF NOT EXISTS books (
                id TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                source_type TEXT NOT NULL,
                source_filename TEXT NOT NULL,
                source_path TEXT NOT NULL,
                extracted_path TEXT NOT NULL,
                status TEXT NOT NULL,
                draft_annotation_path TEXT,
                approved_annotation_path TEXT,
                chapter_count INTEGER NOT NULL DEFAULT 0,
                paragraph_count INTEGER NOT NULL DEFAULT 0,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            """
        )


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def sanitize_stem(name: str) -> str:
    stem = Path(name).stem.strip() or "book"
    sanitized = re.sub(r"[^a-zA-Z0-9_-]+", "-", stem).strip("-")
    return sanitized or "book"


def build_annotated_download_name(book_row: sqlite3.Row) -> str:
    source_filename = book_row["source_filename"] or "book.docx"
    source_stem = Path(source_filename).stem.strip() or "book"
    safe_stem = re.sub(r'[\\/:*?"<>|]+', "-", source_stem).strip(" .") or "book"
    return f"{safe_stem}-annotated.docx"


def save_upload(file: UploadFile, allowed_extensions: set[str], destination_dir: Path) -> tuple[Path, str]:
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in allowed_extensions:
        allowed = ", ".join(sorted(allowed_extensions))
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Allowed: {allowed}")

    base_name = sanitize_stem(file.filename or "book")
    destination = destination_dir / f"{uuid.uuid4().hex}-{base_name}{suffix}"
    with destination.open("wb") as target:
        shutil.copyfileobj(file.file, target)
    return destination, suffix


def extract_docx(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs: list[dict[str, str]] = []
    document_title: str | None = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if style_name.startswith("title"):
            if document_title is None:
                document_title = text
            continue
        if style_name.startswith("subtitle"):
            continue
        paragraphs.append(
            {
                "text": text,
                "style": style_name,
            }
        )
    if not paragraphs:
        raise HTTPException(status_code=400, detail="DOCX file did not contain extractable paragraphs.")
    return {
        "document_title": document_title,
        "paragraphs": paragraphs,
    }


def extract_pdf(path: Path) -> dict[str, Any]:
    reader = PdfReader(str(path))
    paragraphs: list[dict[str, str]] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for block in re.split(r"\n\s*\n", text):
            cleaned = " ".join(line.strip() for line in block.splitlines() if line.strip()).strip()
            if cleaned:
                paragraphs.append(
                    {
                        "text": cleaned,
                        "style": "",
                    }
                )
    if not paragraphs:
        raise HTTPException(status_code=400, detail="PDF file did not contain extractable text.")
    return {
        "document_title": None,
        "paragraphs": paragraphs,
    }


def is_chapter_heading(paragraph: dict[str, str]) -> bool:
    text = paragraph["text"].strip()
    style_name = paragraph.get("style", "")
    if style_name.startswith("heading 1"):
        return True
    if style_name:
        return False
    return bool(CHAPTER_PATTERN.match(text))


def split_into_chapters(paragraphs: list[dict[str, str]]) -> list[dict[str, Any]]:
    chapters: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    opening_paragraphs: list[str] = []

    for paragraph in paragraphs:
        text = paragraph["text"].strip()
        if not text:
            continue

        if is_chapter_heading(paragraph):
            if current is not None and current["paragraphs"]:
                chapters.append(current)
            if current is None and opening_paragraphs:
                current = {
                    "title": text,
                    "leading_paragraphs": opening_paragraphs,
                    "paragraphs": [],
                }
                opening_paragraphs = []
            else:
                current = {
                    "title": text,
                    "leading_paragraphs": [],
                    "paragraphs": [],
                }
            continue

        if current is None:
            opening_paragraphs.append(text)
            continue

        current["paragraphs"].append(text)

    if current is not None and current["paragraphs"]:
        chapters.append(current)

    if not chapters:
        fallback_paragraphs = opening_paragraphs or [paragraph["text"] for paragraph in paragraphs if paragraph["text"].strip()]
        chapters = [{"title": "Chapter 1", "leading_paragraphs": [], "paragraphs": fallback_paragraphs}]

    return chapters


def build_extracted_payload(
    book_id: str,
    title: str,
    source_type: str,
    document_title: str | None,
    paragraphs: list[dict[str, str]],
) -> dict[str, Any]:
    chapters = split_into_chapters(paragraphs)
    paragraph_texts = [paragraph["text"] for paragraph in paragraphs]
    return {
        "book_id": book_id,
        "title": title,
        "document_title": document_title,
        "source_type": source_type,
        "paragraph_count": len(paragraph_texts),
        "chapter_count": len(chapters),
        "chapters": chapters,
    }


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def build_preview_text(extracted: dict[str, Any], limit: int = 2800) -> str:
    preview_parts: list[str] = []
    for chapter in extracted["chapters"]:
        preview_parts.extend(chapter.get("leading_paragraphs", [])[:3])
        preview_parts.append(f"# {chapter['title']}")
        preview_parts.extend(chapter["paragraphs"][:3])
        if len("\n\n".join(preview_parts)) >= limit:
            break
    preview = "\n\n".join(preview_parts)
    return preview[:limit].strip()


def get_gemini_api_key() -> str | None:
    return os.getenv("GEMINI_API_KEY")


def get_gemini_text_model() -> str:
    return os.getenv("GEMINI_TEXT_MODEL", DEFAULT_GEMINI_TEXT_MODEL)


def get_gemini_tts_model() -> str:
    return os.getenv("GEMINI_TTS_MODEL", DEFAULT_GEMINI_TTS_MODEL)


def get_gemini_client() -> Any:
    api_key = get_gemini_api_key()
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY is not configured on the server.")
    try:
        from google import genai
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="The google-genai package is not installed.") from exc
    return genai.Client(api_key=api_key)


def get_gemini_types() -> Any:
    try:
        from google.genai import types
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="The google-genai package is not installed.") from exc
    return types


def infer_segment_type(paragraph: str) -> SegmentType:
    stripped = paragraph.strip()
    if stripped.startswith('"') or stripped.startswith("'"):
        return "dialogue"
    if len(stripped) < 90 and stripped.endswith(":"):
        return "heading"
    return "narration"


def infer_style(segment_type: SegmentType, paragraph: str) -> StyleTag:
    lowered = paragraph.lower()
    if segment_type == "dialogue":
        if any(token in lowered for token in ["!", "run", "quick", "hurry"]):
            return "excited"
        return "neutral"
    return "neutral"


def infer_pace(paragraph: str) -> PaceTag:
    if len(paragraph) > 320:
        return "slow"
    if any(token in paragraph for token in ["!", "?"]):
        return "fast"
    return "steady"


def infer_pause_after(paragraph: str) -> PauseTag:
    if paragraph.endswith((".", "!", "?")):
        return "short"
    return "none"


def build_annotations_from_extracted(extracted: dict[str, Any]) -> AnnotationDocument:
    chapters: list[AnnotatedChapter] = []
    segment_counter = 1
    for chapter_index, chapter in enumerate(extracted["chapters"]):
        segments: list[SegmentAnnotation] = []
        leading_paragraphs = chapter.get("leading_paragraphs", [])
        ordered_paragraphs = [*leading_paragraphs, *chapter["paragraphs"]]
        for paragraph in ordered_paragraphs:
            segment_type = infer_segment_type(paragraph)
            segments.append(
                SegmentAnnotation(
                    segment_id=f"seg_{segment_counter:05d}",
                    segment_type=segment_type,
                    speak=True,
                    clean_text=paragraph,
                    style=infer_style(segment_type, paragraph),
                    pace=infer_pace(paragraph),
                    pause_after=infer_pause_after(paragraph),
                    speaker_hint=None,
                    notes=None,
                )
            )
            segment_counter += 1
        chapters.append(
            AnnotatedChapter(
                chapter_id=f"ch_{chapter_index + 1:03d}",
                chapter_title=chapter["title"],
                leading_segment_count=len(leading_paragraphs),
                segments=segments,
            )
        )
    return AnnotationDocument(
        title=extracted.get("document_title") or "",
        voice_name=DEFAULT_VOICE_NAME,
        chapters=chapters,
    )


def validate_annotated_chapter(
    source_chapter: AnnotatedChapter,
    annotated_chapter: AnnotatedChapter,
    expressive_mode: TagGenerationMode,
    vocalization_mode: TagGenerationMode,
) -> AnnotatedChapter:
    expected_segment_ids = [segment.segment_id for segment in source_chapter.segments]
    actual_segment_ids = [segment.segment_id for segment in annotated_chapter.segments]
    if expected_segment_ids != actual_segment_ids:
        raise HTTPException(
            status_code=502,
            detail="Gemini annotation output did not preserve segment ids and ordering for the chapter.",
        )
    sanitized_segments: list[SegmentAnnotation] = []
    for source_segment, annotated_segment in zip(source_chapter.segments, annotated_chapter.segments, strict=True):
        inferred_style = infer_supported_style(source_segment)
        resolved_style = annotated_segment.style
        if resolved_style == "neutral" and inferred_style != "neutral":
            resolved_style = inferred_style
        elif not style_is_supported_by_text(source_segment, resolved_style):
            resolved_style = inferred_style

        resolved_pace = annotated_segment.pace
        if annotated_segment.pace != "steady" and infer_supported_pace(source_segment) == "steady":
            resolved_pace = "steady"

        resolved_inline_tags = sanitize_generated_inline_tags(
            source_segment,
            annotated_segment.inline_tags,
            expressive_mode,
            vocalization_mode,
        )

        sanitized_segments.append(
            SegmentAnnotation(
                segment_id=source_segment.segment_id,
                segment_type=annotated_segment.segment_type,
                speak=True,
                clean_text=source_segment.clean_text,
                style=resolved_style,
                pace=resolved_pace,
                pause_after=annotated_segment.pause_after,
                inline_tags=resolved_inline_tags,
                speaker_hint=annotated_segment.speaker_hint,
                notes=annotated_segment.notes,
            )
        )
    return AnnotatedChapter(
        chapter_id=source_chapter.chapter_id,
        chapter_title=source_chapter.chapter_title,
        leading_segment_count=source_chapter.leading_segment_count,
        segments=sanitized_segments,
    )


def split_chapter_for_annotation(
    chapter: AnnotatedChapter,
    max_segments: int = DEFAULT_ANNOTATION_BATCH_MAX_SEGMENTS,
    max_chars: int = DEFAULT_ANNOTATION_BATCH_MAX_CHARS,
) -> list[AnnotatedChapter]:
    batches: list[AnnotatedChapter] = []
    batch_segments: list[SegmentAnnotation] = []
    batch_chars = 0

    def flush_batch() -> None:
        nonlocal batch_segments, batch_chars
        if not batch_segments:
            return
        batches.append(
            AnnotatedChapter(
                chapter_id=chapter.chapter_id,
                chapter_title=chapter.chapter_title,
                leading_segment_count=0,
                segments=batch_segments,
            )
        )
        batch_segments = []
        batch_chars = 0

    for segment in chapter.segments:
        segment_chars = len(segment.clean_text)
        would_exceed_segments = len(batch_segments) >= max_segments
        would_exceed_chars = batch_segments and batch_chars + segment_chars > max_chars
        if would_exceed_segments or would_exceed_chars:
            flush_batch()
        batch_segments.append(segment)
        batch_chars += segment_chars

    flush_batch()
    return batches


def annotate_chapter_with_gemini(
    book_title: str,
    chapter: AnnotatedChapter,
    expressive_mode: TagGenerationMode = DEFAULT_EXPRESSIVE_TAG_MODE,
    vocalization_mode: TagGenerationMode = DEFAULT_VOCALIZATION_TAG_MODE,
) -> AnnotatedChapter:
    client = get_gemini_client()
    source_segments = [
        {"segment_id": segment.segment_id, "text": segment.clean_text}
        for segment in chapter.segments
    ]
    prompt = (
        f"Annotate this chapter from the book '{book_title}' for audiobook narration.\n"
        "Return JSON that matches the provided schema exactly.\n"
        "Keep the chapter_id, chapter_title, segment ids, and segment order unchanged.\n"
        "Add annotations only. Do not rewrite, shorten, expand, normalize, or otherwise alter the source text.\n"
        "For every segment, copy the source text into clean_text exactly as provided.\n"
        "For inline_tags, return an empty list unless a tag is clearly justified. Use at most one inline tag total per segment. Do not stack multiple inline tags on the same segment.\n"
        "Do not split, merge, drop, or reorder segments.\n"
        "Do not invent new story content.\n"
        "Use only these enums:\n"
        "segment_type: heading, narration, dialogue, quote, front_matter, back_matter, other\n"
        "style: neutral, serious, warm, curious, tense, sad, excited, soft, whisper, shout, tired, amazed, trembling, panicked, sarcastic, reluctant\n"
        "pace: slow, steady, fast\n"
        "pause_after: none, short, medium, long\n"
        f"Expressive delivery mode: {expressive_mode}.\n"
        f"Vocalization mode: {vocalization_mode}.\n"
        f"Allowed expressive delivery tags for this mode: {', '.join(allowed_inline_tags_for_mode(expressive_mode, expressive_category=True)) or '(none)'}.\n"
        f"Allowed vocalization tags for this mode: {', '.join(allowed_inline_tags_for_mode(vocalization_mode, expressive_category=False)) or '(none)'}.\n"
        "Be conservative. Most segments should remain neutral, steady, and none.\n"
        "Only use expressive style or pace for spoken dialogue or clearly marked internal thoughts.\n"
        "Only use inline_tags for spoken dialogue or clearly marked internal thoughts. Keep ordinary narration plain.\n"
        "Do not add expressive style to ordinary narration.\n"
        "Prefer style over pace. Do not stack style and pace on the same segment.\n"
        "Do not use pause_after to sprinkle [short pause] through ordinary text. Reserve pauses for clear rhetorical breaks, interruptions, or intentionally delayed lines.\n"
        "If a mode is off, do not generate tags from that category.\n"
        "At conservative mode, add a tag only when the wording directly states it or it is extremely obvious.\n"
        "At balanced mode, add a tag when the wording states it or strongly implies it.\n"
        "At expressive mode, add a tag when it is stated, strongly implied, or a clear performance fit.\n"
        "Respect the allowed tag lists for the active modes. Do not use tags outside those lists.\n"
        "Use shout only when the text itself clearly indicates shouting, yelling, screaming, strong all-caps, or repeated exclamation.\n"
        "Use whisper only when the text clearly indicates whispering, muttering, murmuring, or speaking under the breath.\n"
        "Use curious mainly for exploratory spoken questions. Use excited mainly for spoken exclamations.\n"
        "Use tired for weary or drained speech, amazed for shock or discovery, trembling for fearful unsteady speech, panicked for urgent fearful speech, sarcastic for obvious irony, and reluctant for unwilling or hesitant responses.\n"
        "If the text does not support a specific delivery cue, keep the style neutral.\n"
        "Set speak=false for clearly non-spoken material such as title-page remnants, decorative headings, repeated book titles, or metadata.\n"
        "Do not classify normal chapter body text as front_matter just because it appears near the beginning of the file.\n\n"
        f"Chapter id: {chapter.chapter_id}\n"
        f"Chapter title: {chapter.chapter_title}\n\n"
        "Source segments JSON:\n"
        f"{json.dumps(source_segments, ensure_ascii=False, indent=2)}"
    )
    response = client.models.generate_content(
        model=get_gemini_text_model(),
        contents=prompt,
        config={
            "system_instruction": (
                "You annotate book text for audiobook generation. Add metadata only, preserve the source text exactly, "
                "preserve ids, preserve order, and return only valid JSON matching the schema."
            ),
            "response_mime_type": "application/json",
            "response_schema": AnnotatedChapter,
        },
    )
    if not getattr(response, "text", None):
        raise HTTPException(status_code=502, detail="Gemini annotation call did not return structured text.")
    annotated = AnnotatedChapter.model_validate_json(response.text)
    return validate_annotated_chapter(chapter, annotated, expressive_mode, vocalization_mode)


def annotate_chapter_in_batches_with_gemini(
    book_title: str,
    chapter: AnnotatedChapter,
    expressive_mode: TagGenerationMode = DEFAULT_EXPRESSIVE_TAG_MODE,
    vocalization_mode: TagGenerationMode = DEFAULT_VOCALIZATION_TAG_MODE,
) -> AnnotatedChapter:
    batches = split_chapter_for_annotation(chapter)
    if len(batches) <= 1:
        return annotate_chapter_with_gemini(book_title, chapter, expressive_mode, vocalization_mode)

    annotated_segments: list[SegmentAnnotation] = []
    for batch_index, batch in enumerate(batches, start=1):
        try:
            annotated_batch = annotate_chapter_with_gemini(
                book_title,
                batch,
                expressive_mode,
                vocalization_mode,
            )
        except Exception as exc:
            detail = exc.detail if isinstance(exc, HTTPException) else str(exc)
            raise HTTPException(
                status_code=502,
                detail=(
                    f"Gemini annotation failed for chapter '{chapter.chapter_title}', "
                    f"batch {batch_index} of {len(batches)}: {detail}"
                ),
            ) from exc
        annotated_segments.extend(annotated_batch.segments)

    return AnnotatedChapter(
        chapter_id=chapter.chapter_id,
        chapter_title=chapter.chapter_title,
        leading_segment_count=chapter.leading_segment_count,
        segments=annotated_segments,
    )


def generate_annotation_document(
    extracted: dict[str, Any],
    expressive_mode: TagGenerationMode = DEFAULT_EXPRESSIVE_TAG_MODE,
    vocalization_mode: TagGenerationMode = DEFAULT_VOCALIZATION_TAG_MODE,
    progress_callback: Callable[[str, int, int, str | None], None] | None = None,
) -> tuple[AnnotationDocument, str]:
    heuristic_annotations = build_annotations_from_extracted(extracted)
    total_chapters = len(heuristic_annotations.chapters)
    if progress_callback is not None:
        progress_callback("Preparing chapters", 0, total_chapters, None)
    if not get_gemini_api_key():
        if progress_callback is not None:
            progress_callback("Writing DOCX", total_chapters, total_chapters, None)
        return heuristic_annotations, "heuristic"
    annotated_chapters: list[AnnotatedChapter] = []
    for chapter_number, chapter in enumerate(heuristic_annotations.chapters, start=1):
        if progress_callback is not None:
            progress_callback("Annotating chapters", chapter_number - 1, total_chapters, chapter.chapter_title)
        annotated_chapters.append(
            annotate_chapter_in_batches_with_gemini(
                heuristic_annotations.title,
                chapter,
                expressive_mode,
                vocalization_mode,
            )
        )
        if progress_callback is not None:
            progress_callback("Annotating chapters", chapter_number, total_chapters, chapter.chapter_title)
    if progress_callback is not None:
        progress_callback("Writing DOCX", total_chapters, total_chapters, None)
    return AnnotationDocument(
        title=heuristic_annotations.title,
        voice_name=heuristic_annotations.voice_name,
        chapters=annotated_chapters,
    ), "gemini"


def create_draft_annotation_docx(
    book_row: sqlite3.Row,
    annotations: AnnotationDocument,
    expressive_mode: TagGenerationMode = DEFAULT_EXPRESSIVE_TAG_MODE,
    vocalization_mode: TagGenerationMode = DEFAULT_VOCALIZATION_TAG_MODE,
) -> Path:
    document = Document()
    for chapter in annotations.chapters:
        leading_count = chapter.leading_segment_count

        for segment in chapter.segments[:leading_count]:
            document.add_paragraph(build_inline_tts_text(segment, expressive_mode, vocalization_mode))

        document.add_heading(chapter.chapter_title, level=1)
        first_spoken_paragraph_written = False
        for segment in chapter.segments[leading_count:]:
            paragraph_text = build_inline_tts_text(segment, expressive_mode, vocalization_mode)
            if segment.speak and not first_spoken_paragraph_written:
                if not paragraph_text.strip().startswith("[short pause]"):
                    paragraph_text = f"[short pause] {paragraph_text.strip()}".strip()
                first_spoken_paragraph_written = True
            document.add_paragraph(paragraph_text)

    path = ANNOTATED_DIR / f"{book_row['id']}-draft.docx"
    document.save(path)
    return path


def parse_annotation_docx(path: Path) -> AnnotationDocument:
    document = Document(path)
    chapters: list[AnnotatedChapter] = []
    pending_segments: list[SegmentAnnotation] = []
    current_title: str | None = None
    current_segments: list[SegmentAnnotation] = []
    current_leading_segment_count = 0
    title: str | None = None
    segment_counter = 1

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if style_name.startswith("heading 1"):
            if current_title is not None:
                chapters.append(
                    AnnotatedChapter(
                        chapter_id=f"ch_{len(chapters) + 1:03d}",
                        chapter_title=current_title,
                        leading_segment_count=current_leading_segment_count,
                        segments=current_segments,
                    )
                )
            current_title = text
            current_segments = pending_segments if pending_segments else []
            current_leading_segment_count = len(pending_segments)
            pending_segments = []
            continue
        if style_name.startswith("title"):
            title = text
            continue
        validate_inline_tts_tags(text)
        bare_text = strip_inline_tts_tags(text)
        inferred_segment_type = infer_segment_type(bare_text or text)
        segment = SegmentAnnotation(
            segment_id=f"seg_{segment_counter:05d}",
            segment_type=inferred_segment_type,
            speak=True,
            clean_text=text,
            style="neutral",
            pace="steady",
            pause_after="none",
            inline_tags=extract_inline_tts_tags(text),
            speaker_hint=None,
            notes=None,
        )
        if current_title is None:
            pending_segments.append(segment)
        else:
            current_segments.append(segment)
        segment_counter += 1

    if current_title is None and pending_segments:
        chapters.append(
            AnnotatedChapter(
                chapter_id="ch_001",
                chapter_title="",
                leading_segment_count=0,
                segments=pending_segments,
            )
        )
    elif current_title is not None:
        chapters.append(
            AnnotatedChapter(
                chapter_id=f"ch_{len(chapters) + 1:03d}",
                chapter_title=current_title,
                leading_segment_count=current_leading_segment_count,
                segments=current_segments,
            )
        )

    if not chapters:
        raise HTTPException(status_code=400, detail="Approved DOCX did not contain chapter content.")

    return AnnotationDocument(title=title or path.stem, voice_name=DEFAULT_VOICE_NAME, chapters=chapters)


def chunk_chapters(
    chapters: list[AnnotatedChapter],
    max_chars: int,
    hard_max_chars: int | None = None,
    selected_chapter_indexes: list[int] | None = None,
) -> list[dict[str, Any]]:
    max_chars, hard_max_chars = normalize_chunk_limits(max_chars, hard_max_chars)

    def flush_chunk(
        items: list[SegmentAnnotation],
        chapter_index: int,
        chapter_title: str,
        chunk_index: int,
    ) -> dict[str, Any]:
        joined = build_chunk_text(items)
        paragraph_count = sum(1 for item in items if item.clean_text.strip() and not item.join_with_previous)
        return {
            "chapter_index": chapter_index,
            "chapter_title": chapter_title,
            "chunk_index": chunk_index,
            "char_count": len(joined),
            "paragraph_count": paragraph_count,
            "text": joined,
            "preview": joined[:220],
            "segments": [item.model_dump() for item in items],
        }

    chunks: list[dict[str, Any]] = []
    selected_indexes = set(selected_chapter_indexes) if selected_chapter_indexes is not None else None
    for chapter_index, chapter in enumerate(chapters):
        if selected_indexes is not None and chapter_index not in selected_indexes:
            continue
        bucket: list[SegmentAnnotation] = []
        bucket_chars = 0
        chunk_index = 0
        for segment in chapter.segments:
            if not segment.speak:
                continue
            normalized_segment = segment.model_copy(update={"clean_text": segment.clean_text.strip()})
            if not normalized_segment.clean_text:
                continue
            if bucket and bucket_chars >= max_chars:
                chunks.append(flush_chunk(bucket, chapter_index, chapter.chapter_title, chunk_index))
                chunk_index += 1
                bucket = []
                bucket_chars = 0
            if bucket_chars + measure_segment_size(normalized_segment, bool(bucket)) <= hard_max_chars:
                bucket.append(normalize_chunk_segment(normalized_segment, bool(bucket)))
                bucket_chars = len(build_chunk_text(bucket))
                continue

            for split_segment in split_segment_for_chunking(normalized_segment, hard_max_chars):
                prepared_segment = normalize_chunk_segment(split_segment, bool(bucket))
                segment_size = measure_segment_size(prepared_segment, bool(bucket))
                if bucket and bucket_chars + segment_size > hard_max_chars:
                    chunks.append(flush_chunk(bucket, chapter_index, chapter.chapter_title, chunk_index))
                    chunk_index += 1
                    bucket = []
                    bucket_chars = 0
                    prepared_segment = normalize_chunk_segment(split_segment, False)
                    segment_size = measure_segment_size(prepared_segment, False)
                bucket.append(prepared_segment)
                bucket_chars += segment_size
        if bucket:
            chunks.append(flush_chunk(bucket, chapter_index, chapter.chapter_title, chunk_index))
    return chunks


def normalize_chunk_limits(max_chars: int, hard_max_chars: int | None = None) -> tuple[int, int]:
    target = int(max_chars or DEFAULT_MAX_CHARS)
    hard_limit = int(hard_max_chars or target + 100)
    if target < 100 or target > CHUNK_TARGET_PRESETS[-1]:
        raise HTTPException(status_code=400, detail="Chunk limit must be between 100 and 1500 characters.")
    if hard_limit < target:
        raise HTTPException(status_code=400, detail="Hard limit must be greater than or equal to the chunk limit.")
    if hard_limit > MAX_HARD_MAX_CHARS:
        raise HTTPException(status_code=400, detail=f"Hard limit must be {MAX_HARD_MAX_CHARS} characters or less.")
    return target, hard_limit


def normalize_chunk_segment(segment: SegmentAnnotation, has_existing_bucket: bool) -> SegmentAnnotation:
    if has_existing_bucket:
        return segment
    if not segment.join_with_previous:
        return segment
    return segment.model_copy(update={"join_with_previous": False})


def build_chunk_text(segments: list[SegmentAnnotation]) -> str:
    paragraphs: list[str] = []
    for segment in segments:
        clean_text = segment.clean_text.strip()
        if not clean_text:
            continue
        if paragraphs and segment.join_with_previous:
            paragraphs[-1] = f"{paragraphs[-1]} {clean_text}"
            continue
        paragraphs.append(clean_text)
    return "\n\n".join(paragraphs)


def measure_segment_size(segment: SegmentAnnotation, has_existing_bucket: bool) -> int:
    clean_text = segment.clean_text.strip()
    if not clean_text:
        return 0
    return len(clean_text) + (1 if has_existing_bucket else 0)


def split_segment_for_chunking(segment: SegmentAnnotation, hard_max_chars: int) -> list[SegmentAnnotation]:
    text = segment.clean_text.strip()
    if len(text) <= hard_max_chars:
        return [segment.model_copy(update={"clean_text": text})]

    sentence_units = split_text_to_units(text, hard_max_chars, SENTENCE_BOUNDARY_PATTERN)
    split_segments: list[SegmentAnnotation] = []
    for unit_index, unit in enumerate(sentence_units):
        split_segments.append(
            segment.model_copy(
                update={
                    "segment_id": f"{segment.segment_id}_part_{unit_index + 1}",
                    "clean_text": unit,
                    "join_with_previous": unit_index > 0,
                }
            )
        )
    return split_segments


def split_text_to_units(text: str, hard_max_chars: int, boundary_pattern: re.Pattern[str]) -> list[str]:
    pieces = [piece.strip() for piece in boundary_pattern.split(text.strip()) if piece.strip()]
    if len(pieces) <= 1:
        if boundary_pattern is SENTENCE_BOUNDARY_PATTERN:
            return split_text_to_units(text, hard_max_chars, CLAUSE_BOUNDARY_PATTERN)
        return split_text_by_length(text, hard_max_chars)

    grouped_units: list[str] = []
    bucket = ""
    for piece in pieces:
        candidate = piece if not bucket else f"{bucket} {piece}"
        if bucket and len(candidate) > hard_max_chars:
            grouped_units.append(bucket)
            bucket = piece
            continue
        bucket = candidate
    if bucket:
        grouped_units.append(bucket)

    units: list[str] = []
    for grouped_unit in grouped_units:
        if len(grouped_unit) > hard_max_chars:
            units.extend(split_text_by_length(grouped_unit, hard_max_chars))
        else:
            units.append(grouped_unit)
    return units


def split_text_by_length(text: str, hard_max_chars: int) -> list[str]:
    remaining = text.strip()
    parts: list[str] = []
    while len(remaining) > hard_max_chars:
        split_at = remaining.rfind(" ", 0, hard_max_chars + 1)
        if split_at <= 0:
            split_at = hard_max_chars
        parts.append(remaining[:split_at].strip())
        remaining = remaining[split_at:].strip()
    if remaining:
        parts.append(remaining)
    return parts


def get_annotation_document(extracted: dict[str, Any]) -> AnnotationDocument:
    approved = extracted.get("approved_annotation")
    if approved:
        return AnnotationDocument.model_validate(approved)
    generated = extracted.get("generated_annotation")
    if generated:
        return AnnotationDocument.model_validate(generated)
    return build_annotations_from_extracted(extracted)


def build_extracted_payload_from_annotation_document(
    book_id: str,
    annotation_document: AnnotationDocument,
    source_filename: str,
) -> dict[str, Any]:
    chapters: list[dict[str, Any]] = []
    paragraph_count = 0
    for chapter in annotation_document.chapters:
        paragraphs = [segment.clean_text for segment in chapter.segments if segment.clean_text.strip()]
        leading_count = min(chapter.leading_segment_count, len(paragraphs))
        leading_paragraphs = paragraphs[:leading_count]
        chapter_paragraphs = paragraphs[leading_count:]
        paragraph_count += len(paragraphs)
        chapters.append(
            {
                "title": chapter.chapter_title,
                "leading_paragraphs": leading_paragraphs,
                "paragraphs": chapter_paragraphs,
            }
        )

    return {
        "book_id": book_id,
        "title": annotation_document.title or Path(source_filename).stem,
        "document_title": annotation_document.title or Path(source_filename).stem,
        "source_type": "docx",
        "paragraph_count": paragraph_count,
        "chapter_count": len(chapters),
        "chapters": chapters,
        "approved_annotation": annotation_document.model_dump(),
    }


def render_segment_for_tts(segment: SegmentAnnotation) -> str:
    if extract_inline_tts_tags(segment.clean_text):
        validate_inline_tts_tags(segment.clean_text)
        return segment.clean_text.strip()
    return build_inline_tts_text(segment)


def build_gemini_tts_prompt(
    book_title: str,
    chapter_title: str,
    voice_name: str,
    language_code: str,
    style_instruction: str,
    segments: list[SegmentAnnotation],
) -> str:
    transcript = build_tts_transcript(segments)
    chapter_context = f"Chapter context: {chapter_title}.\n" if chapter_title.strip() else ""
    return (
        "Please synthesize speech audio for the transcript below.\n"
        "Read only the transcript. Do not read headings, notes, or metadata aloud.\n\n"
        f"### BOOK\n{book_title}\n\n"
        "### DIRECTOR'S NOTES\n"
        f"{chapter_context}"
        f"Voice: {voice_name}.\n"
        f"Language: {language_code}.\n"
        f"Style: {style_instruction}\n"
        "Pacing: Respect paragraph breaks and inline pause tags.\n"
        "Continuity: Use one continuous narrator voice across the entire transcript. Do not switch to a different voice for pre-chapter text, chapter titles, or the opening body paragraphs unless inline tags explicitly call for it.\n"
        "Accent: Neutral modern spoken English unless the transcript strongly suggests otherwise.\n\n"
        "### TRANSCRIPT\n"
        f"{transcript}"
    )


def build_tts_transcript(segments: list[SegmentAnnotation]) -> str:
    transcript_parts: list[str] = []
    for segment in segments:
        if not segment.speak:
            continue
        rendered_segment = render_segment_for_tts(segment)
        if transcript_parts and segment.join_with_previous:
            transcript_parts[-1] = f"{transcript_parts[-1]} {rendered_segment}"
            continue
        transcript_parts.append(rendered_segment)
    return "\n".join(transcript_parts)


def build_chapter_title_segment(chapter: AnnotatedChapter) -> SegmentAnnotation:
    return SegmentAnnotation(
        segment_id=f"{chapter.chapter_id}_title",
        segment_type="heading",
        speak=True,
        clean_text=chapter.chapter_title,
        style="neutral",
        pace="steady",
        pause_after="none",
        speaker_hint=None,
        notes=None,
    )


def prepare_chapters_for_generation(
    chapters: list[AnnotatedChapter],
    read_chapter_titles: bool,
    pre_chapter_text_mode: PreChapterTextMode,
) -> list[AnnotatedChapter]:
    prepared_chapters: list[AnnotatedChapter] = []
    for chapter_index, chapter in enumerate(chapters):
        chapter_segments = [segment.model_copy() for segment in chapter.segments]
        leading_segment_count = chapter.leading_segment_count

        if chapter_index == 0 and pre_chapter_text_mode == "separate_chapter_0" and leading_segment_count > 0:
            leading_segments = [segment.model_copy(update={"join_with_previous": False}) for segment in chapter_segments[:leading_segment_count]]
            if leading_segments:
                synthetic_chapter_zero = AnnotatedChapter(
                    chapter_id="ch_000",
                    chapter_title="Chapter 0",
                    leading_segment_count=0,
                    segments=leading_segments,
                )
                if read_chapter_titles and synthetic_chapter_zero.chapter_title.strip():
                    synthetic_chapter_zero.segments = [build_chapter_title_segment(synthetic_chapter_zero), *synthetic_chapter_zero.segments]
                prepared_chapters.append(synthetic_chapter_zero)
            chapter_segments = [segment.model_copy(update={"join_with_previous": False}) for segment in chapter_segments[leading_segment_count:]]
            leading_segment_count = 0

        prepared_chapter = AnnotatedChapter(
            chapter_id=chapter.chapter_id,
            chapter_title=chapter.chapter_title,
            leading_segment_count=leading_segment_count,
            segments=chapter_segments,
        )
        if read_chapter_titles and prepared_chapter.chapter_title.strip():
            title_segment = build_chapter_title_segment(prepared_chapter)
            if chapter_index == 0 and pre_chapter_text_mode == "attach_to_chapter_1" and leading_segment_count > 0:
                prepared_chapter.segments = [
                    *prepared_chapter.segments[:leading_segment_count],
                    title_segment,
                    *prepared_chapter.segments[leading_segment_count:],
                ]
            else:
                prepared_chapter.segments = [title_segment, *prepared_chapter.segments]
            prepared_chapter.leading_segment_count = 0
        prepared_chapters.append(prepared_chapter)

    return [chapter for chapter in prepared_chapters if chapter.segments]


def normalize_transcript_for_audit(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def find_missing_segment_excerpts(
    source_segments: list[SegmentAnnotation],
    prepared_transcript: str,
    limit: int = 5,
) -> list[str]:
    prepared_normalized = normalize_transcript_for_audit(prepared_transcript)
    search_start = 0
    missing_excerpts: list[str] = []
    for segment in source_segments:
        if not segment.speak:
            continue
        source_text = normalize_transcript_for_audit(render_segment_for_tts(segment))
        if not source_text:
            continue
        position = prepared_normalized.find(source_text, search_start)
        if position == -1:
            excerpt = segment.clean_text.strip().replace("\n", " ")
            if excerpt and excerpt not in missing_excerpts:
                missing_excerpts.append(excerpt[:180])
            if len(missing_excerpts) >= limit:
                break
            continue
        search_start = position + len(source_text)
    return missing_excerpts


def review_transcript_coverage_with_gemini(
    book_title: str,
    chapter_title: str,
    source_transcript: str,
    prepared_transcript: str,
) -> TranscriptCoverageReview | None:
    if not get_gemini_api_key():
        return None

    client = get_gemini_client()
    prompt = (
        "Compare the source transcript with the prepared transcript that will be sent to TTS.\n"
        "Ignore whitespace-only differences caused by chunk boundaries.\n"
        "Mark is_complete true only if every spoken line from the source appears in the prepared transcript in the same order.\n"
        "If anything is missing or reordered, summarize the problem and list the shortest useful missing excerpts.\n\n"
        f"BOOK: {book_title}\n"
        f"CHAPTER: {chapter_title}\n\n"
        "SOURCE TRANSCRIPT:\n"
        f"{source_transcript}\n\n"
        "PREPARED TRANSCRIPT:\n"
        f"{prepared_transcript}"
    )
    try:
        response = client.models.generate_content(
            model=get_gemini_text_model(),
            contents=prompt,
            config={
                "system_instruction": (
                    "You verify transcript coverage before text-to-speech generation. "
                    "Be strict about omissions and order changes, but ignore whitespace-only formatting differences."
                ),
                "response_mime_type": "application/json",
                "response_schema": TranscriptCoverageReview,
            },
        )
        if not getattr(response, "text", None):
            return None
        return TranscriptCoverageReview.model_validate_json(response.text)
    except Exception:
        return None


def audit_prepared_transcript_coverage(
    book_title: str,
    chapters: list[AnnotatedChapter],
    chunks: list[dict[str, Any]],
    selected_chapter_indexes: list[int],
) -> dict[str, Any]:
    prepared_segments_by_chapter: dict[int, list[SegmentAnnotation]] = {index: [] for index in selected_chapter_indexes}
    for chunk in chunks:
        chapter_index = chunk["chapter_index"]
        if chapter_index not in prepared_segments_by_chapter:
            continue
        prepared_segments_by_chapter[chapter_index].extend(
            SegmentAnnotation.model_validate(segment) for segment in chunk["segments"]
        )

    chapter_reviews: list[dict[str, Any]] = []
    issues: list[dict[str, Any]] = []
    for chapter_index in selected_chapter_indexes:
        chapter = chapters[chapter_index]
        source_transcript = build_tts_transcript(chapter.segments)
        prepared_transcript = build_tts_transcript(prepared_segments_by_chapter.get(chapter_index, []))
        source_normalized = normalize_transcript_for_audit(source_transcript)
        prepared_normalized = normalize_transcript_for_audit(prepared_transcript)
        chapter_ok = source_normalized == prepared_normalized
        missing_excerpts = [] if chapter_ok else find_missing_segment_excerpts(chapter.segments, prepared_transcript)
        review_summary = "Transcript coverage verified."
        llm_review = None if chapter_ok else review_transcript_coverage_with_gemini(
            book_title,
            chapter.chapter_title,
            source_transcript,
            prepared_transcript,
        )
        if llm_review is not None:
            review_summary = llm_review.summary
            if llm_review.missing_excerpts:
                missing_excerpts = llm_review.missing_excerpts[:5]
        elif not chapter_ok:
            review_summary = "Prepared transcript does not match the approved chapter transcript."

        chapter_review = {
            "chapter_index": chapter_index,
            "chapter_title": chapter.chapter_title,
            "is_complete": chapter_ok,
            "summary": review_summary,
            "missing_excerpts": missing_excerpts,
        }
        chapter_reviews.append(chapter_review)
        if not chapter_ok:
            issues.append(chapter_review)

    if issues:
        issue_summaries = []
        for issue in issues[:3]:
            excerpt_text = "; ".join(issue["missing_excerpts"][:2]) if issue["missing_excerpts"] else issue["summary"]
            issue_summaries.append(f"{issue['chapter_title']}: {excerpt_text}")
        raise HTTPException(
            status_code=500,
            detail="Transcript coverage check failed before TTS generation. " + " | ".join(issue_summaries),
        )

    return {
        "checked_chapter_count": len(chapter_reviews),
        "summary": f"Verified transcript coverage for {len(chapter_reviews)} chapter(s).",
        "chapters": chapter_reviews,
    }


def write_wave_file(
    path: Path,
    pcm_data: bytes,
    channels: int = 1,
    rate: int = DEFAULT_AUDIO_SAMPLE_RATE,
    sample_width: int = 2,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with wave.open(str(path), "wb") as wav_file:
        wav_file.setnchannels(channels)
        wav_file.setsampwidth(sample_width)
        wav_file.setframerate(rate)
        wav_file.writeframes(pcm_data)


def concatenate_pcm_chunks(chunks: list[bytes]) -> bytes:
    return b"".join(chunks)


def get_inline_audio_sample_rate(inline_data: Any) -> int:
    mime_type = getattr(inline_data, "mime_type", "") or ""
    match = re.search(r"rate=(\d+)", mime_type, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return DEFAULT_AUDIO_SAMPLE_RATE


def resolve_chunk(chunks: list[dict[str, Any]], chapter_index: int, chunk_index: int) -> dict[str, Any]:
    selected_chunk = next(
        (
            chunk
            for chunk in chunks
            if chunk["chapter_index"] == chapter_index and chunk["chunk_index"] == chunk_index
        ),
        None,
    )
    if selected_chunk is None:
        raise HTTPException(status_code=404, detail="Requested chunk was not found in the current chunk plan.")
    return selected_chunk


def synthesize_chunk_audio_with_gemini(
    prompt: str,
    voice_name: str,
    language_code: str,
    max_attempts: int = DEFAULT_TTS_MAX_ATTEMPTS,
    on_retry: Callable[[int, int, str], None] | None = None,
) -> tuple[bytes, int, int]:
    client = get_gemini_client()
    types = get_gemini_types()
    last_error = "Gemini TTS did not return audio data."

    for attempt in range(1, max_attempts + 1):
        try:
            response = client.models.generate_content(
                model=get_gemini_tts_model(),
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_modalities=["AUDIO"],
                    speech_config=types.SpeechConfig(
                        language_code=language_code,
                        voice_config=types.VoiceConfig(
                            prebuilt_voice_config=types.PrebuiltVoiceConfig(voice_name=voice_name)
                        )
                    ),
                ),
            )
            inline_data = response.candidates[0].content.parts[0].inline_data
            audio_data = inline_data.data
            if not audio_data:
                last_error = "Gemini TTS returned empty audio data."
            else:
                return audio_data, get_inline_audio_sample_rate(inline_data), attempt
        except (AttributeError, IndexError, KeyError, TypeError):
            last_error = "Gemini TTS did not return audio data."
        except Exception as exc:
            last_error = str(exc)

        if attempt < max_attempts:
            if on_retry is not None:
                on_retry(attempt + 1, max_attempts, last_error)
            time.sleep(1)

    raise HTTPException(
        status_code=502,
        detail=f"Gemini TTS failed after {max_attempts} attempts. Last error: {last_error}",
    )


def build_chunk_audio_path(book_id: str, chapter_index: int, chunk_index: int, voice_name: str) -> Path:
    safe_voice_name = sanitize_stem(voice_name)
    return AUDIO_DIR / book_id / f"chapter-{chapter_index + 1:03d}-chunk-{chunk_index + 1:03d}-{safe_voice_name}.wav"


def build_merged_audio_path(book_id: str, voice_name: str, speech_rate: float) -> Path:
    safe_voice_name = sanitize_stem(voice_name)
    speed_label = str(speech_rate).replace('.', '-')
    return AUDIO_DIR / book_id / f"audiobook-{safe_voice_name}-speed-{speed_label}.mp3"


def build_raw_merged_audio_path(book_id: str, voice_name: str) -> Path:
    safe_voice_name = sanitize_stem(voice_name)
    return AUDIO_DIR / book_id / f"raw-audiobook-{safe_voice_name}.wav"


def build_chapter_audio_path(book_id: str, chapter_index: int, chapter_title: str, voice_name: str, speech_rate: float) -> Path:
    safe_voice_name = sanitize_stem(voice_name)
    safe_chapter_title = sanitize_stem(chapter_title) or f"chapter-{chapter_index + 1:03d}"
    speed_label = str(speech_rate).replace('.', '-')
    return AUDIO_DIR / book_id / f"chapter-{chapter_index + 1:03d}-{safe_chapter_title}-{safe_voice_name}-speed-{speed_label}.mp3"


def build_chapter_output_path(
    book_id: str,
    chapter_index: int,
    chapter_title: str,
    voice_name: str,
    speech_rate: float,
    output_format: AudioOutputFormat,
) -> Path:
    safe_voice_name = sanitize_stem(voice_name)
    safe_chapter_title = sanitize_stem(chapter_title) or f"chapter-{chapter_index + 1:03d}"
    speed_label = str(speech_rate).replace('.', '-')
    return AUDIO_DIR / book_id / f"chapter-{chapter_index + 1:03d}-{safe_chapter_title}-{safe_voice_name}-speed-{speed_label}.{output_format}"


def build_raw_chapter_audio_path(book_id: str, chapter_index: int, voice_name: str) -> Path:
    safe_voice_name = sanitize_stem(voice_name)
    return AUDIO_DIR / book_id / f"raw-chapter-{chapter_index + 1:03d}-{safe_voice_name}.wav"


def build_audio_zip_path(book_id: str, voice_name: str, speech_rate: float, output_format: AudioOutputFormat) -> Path:
    safe_voice_name = sanitize_stem(voice_name)
    speed_label = str(speech_rate).replace('.', '-')
    return AUDIO_DIR / book_id / f"chapters-{safe_voice_name}-speed-{speed_label}-{output_format}.zip"


def build_audio_status_path(book_id: str) -> Path:
    return AUDIO_DIR / book_id / "generation-status.json"


def build_draft_annotation_status_path(book_id: str) -> Path:
    return ANNOTATED_DIR / f"{book_id}-generation-status.json"


def write_json_atomic(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    temporary_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    os.replace(temporary_path, path)


def read_json_with_retry(path: Path, attempts: int = 3, delay_seconds: float = 0.05) -> dict[str, Any]:
    last_error: Exception | None = None
    for attempt in range(attempts):
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError) as exc:
            last_error = exc
            if attempt < attempts - 1:
                time.sleep(delay_seconds)
    raise last_error or RuntimeError(f"Failed to read JSON file: {path}")


def write_audio_generation_status(book_id: str, payload: dict[str, Any]) -> None:
    path = build_audio_status_path(book_id)
    write_json_atomic(path, payload)


def read_audio_generation_status(book_id: str) -> dict[str, Any] | None:
    path = build_audio_status_path(book_id)
    if not path.exists():
        return None
    return read_json_with_retry(path)


def write_draft_annotation_status(book_id: str, payload: dict[str, Any]) -> None:
    path = build_draft_annotation_status_path(book_id)
    write_json_atomic(path, payload)


def read_draft_annotation_status(book_id: str) -> dict[str, Any] | None:
    path = build_draft_annotation_status_path(book_id)
    if not path.exists():
        return None
    return read_json_with_retry(path)


def materialize_chapter_audio_output(
    book_id: str,
    chapter_index: int,
    chapter_title: str,
    voice_name: str,
    speech_rate: float,
    output_format: AudioOutputFormat,
    audio_quality: AudioQualityPreset,
    chapter_chunks: list[bytes],
    sample_rate: int,
) -> Path:
    raw_chapter_path = build_raw_chapter_audio_path(book_id, chapter_index, voice_name)
    chapter_output_path = build_chapter_output_path(
        book_id,
        chapter_index,
        chapter_title,
        voice_name,
        speech_rate,
        output_format,
    )
    chapter_pcm = concatenate_pcm_chunks(chapter_chunks)
    write_wave_file(raw_chapter_path, chapter_pcm, rate=sample_rate)
    if output_format == "wav":
        return raw_chapter_path
    encode_mp3_with_speed(raw_chapter_path, chapter_output_path, speech_rate, audio_quality)
    return chapter_output_path


def encode_mp3_with_speed(
    source_path: Path,
    destination_path: Path,
    speech_rate: float,
    audio_quality: AudioQualityPreset,
) -> None:
    destination_path.parent.mkdir(parents=True, exist_ok=True)
    filter_chain = f"atempo={speech_rate}" if abs(speech_rate - 1.0) >= 0.001 else "anull"
    quality_value = "0" if audio_quality == "best" else "2"
    try:
        subprocess.run(
            [
                "ffmpeg",
                "-y",
                "-i",
                str(source_path),
                "-filter:a",
                filter_chain,
                "-codec:a",
                "libmp3lame",
                "-q:a",
                quality_value,
                str(destination_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
    except FileNotFoundError as exc:
        raise HTTPException(status_code=500, detail="ffmpeg is required on the server for MP3 export and audio speed adjustment.") from exc
    except subprocess.CalledProcessError as exc:
        raise HTTPException(status_code=500, detail=f"ffmpeg failed while encoding MP3 audio: {exc.stderr.strip()}") from exc


def get_book_or_404(book_id: str) -> sqlite3.Row:
    with get_connection() as connection:
        row = connection.execute("SELECT * FROM books WHERE id = ?", (book_id,)).fetchone()
    if row is None:
        raise HTTPException(status_code=404, detail="Book not found.")
    return row


def serialize_book(row: sqlite3.Row) -> dict[str, Any]:
    extracted = read_json(Path(row["extracted_path"]))
    draft_settings = normalize_generated_annotation_settings(extracted)
    draft_path = Path(row["draft_annotation_path"]) if row["draft_annotation_path"] else None
    approved_path = Path(row["approved_annotation_path"]) if row["approved_annotation_path"] else None
    annotations = get_annotation_document(extracted)
    audio_status = read_audio_generation_status(row["id"])
    draft_annotation_status = read_draft_annotation_status(row["id"])
    return {
        "id": row["id"],
        "title": row["title"],
        "source_type": row["source_type"],
        "source_filename": row["source_filename"],
        "status": row["status"],
        "chapter_count": row["chapter_count"],
        "paragraph_count": row["paragraph_count"],
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
        "preview_text": build_preview_text(extracted),
        "has_draft_annotation": bool(draft_path and draft_path.exists()),
        "has_approved_annotation": bool(approved_path and approved_path.exists()),
        "draft_annotation_url": f"/api/books/{row['id']}/draft-annotation/download" if draft_path and draft_path.exists() else None,
        "draft_annotation_expressive_mode": draft_settings["expressive_mode"],
        "draft_annotation_vocalization_mode": draft_settings["vocalization_mode"],
        "voice_name": annotations.voice_name,
        "generator_settings": get_generator_settings(extracted),
        "chapters": build_chapter_summaries(annotations),
        "annotation_provider": extracted.get("generated_annotation_provider", "heuristic"),
        "draft_annotation_generation": draft_annotation_status,
        "audio_generation": audio_status,
        "annotation_stats": {
            "chapter_count": len(annotations.chapters),
            "segment_count": sum(len(chapter.segments) for chapter in annotations.chapters),
        },
    }


def run_book_audio_generation(book_id: str, request: GenerateBookAudioRequest) -> None:
    generated_chapters: list[dict[str, Any]] = []

    def build_audio_status_payload(**extra: Any) -> dict[str, Any]:
        payload = {
            "download_url": None,
            "file_name": None,
            "speech_rate": request.speech_rate,
            "voice_name": request.voice_name,
            "language_code": request.language_code,
            "style_instruction": request.style_instruction,
            "output_format": request.output_format,
            "audio_quality": request.audio_quality,
            "selected_chapter_indexes": extra.get("selected_chapter_indexes"),
            "generated_chapters": list(generated_chapters),
            "updated_at": now_iso(),
        }
        payload.update(extra)
        return payload

    try:
        book = get_book_or_404(book_id)
        extracted = read_json(Path(book["extracted_path"]))
        approved = extracted.get("approved_annotation")
        if not approved:
            raise HTTPException(status_code=400, detail="Upload an approved annotation DOCX before generating audio.")

        annotations = AnnotationDocument.model_validate(approved)
        prepared_chapters = prepare_chapters_for_generation(
            annotations.chapters,
            request.read_chapter_titles,
            request.pre_chapter_text_mode,
        )
        selected_chapter_indexes = normalize_selected_chapter_indexes(len(prepared_chapters), request.chapter_indexes)
        chunks = chunk_chapters(
            prepared_chapters,
            request.max_chars,
            request.hard_max_chars,
            selected_chapter_indexes,
        )
        if not chunks:
            raise HTTPException(status_code=400, detail="No audio chunks were produced for the current annotation.")

        write_audio_generation_status(
            book_id,
            build_audio_status_payload(
                state="running",
                step="Verifying transcript coverage",
                chunk_index=0,
                chunk_count=len(chunks),
                hard_max_chars=request.hard_max_chars,
                selected_chapter_indexes=selected_chapter_indexes,
            ),
        )
        coverage_audit = audit_prepared_transcript_coverage(
            annotations.title,
            prepared_chapters,
            chunks,
            selected_chapter_indexes,
        )

        write_audio_generation_status(
            book_id,
            build_audio_status_payload(
                state="running",
                step="Preparing audiobook generation",
                chunk_index=0,
                chunk_count=len(chunks),
                hard_max_chars=request.hard_max_chars,
                selected_chapter_indexes=selected_chapter_indexes,
                coverage_check=coverage_audit,
            ),
        )

        chapter_audio: dict[int, list[bytes]] = {}
        chapter_sample_rates: dict[int, int] = {}
        chapter_titles: dict[int, str] = {}
        chapter_output_paths: list[Path] = []
        chapter_chunk_totals: dict[int, int] = {}
        for chunk in chunks:
            chapter_index = chunk["chapter_index"]
            chapter_chunk_totals[chapter_index] = chapter_chunk_totals.get(chapter_index, 0) + 1
        chapter_chunk_positions: dict[int, int] = {}
        for position, chunk in enumerate(chunks, start=1):
            chapter_index = chunk["chapter_index"]
            chapter_chunk_positions[chapter_index] = chapter_chunk_positions.get(chapter_index, 0) + 1
            write_audio_generation_status(
                book_id,
                build_audio_status_payload(
                    state="running",
                    step=f"Generating chunk {position} of {len(chunks)}",
                    chunk_index=position,
                    chunk_count=len(chunks),
                    active_chapter_index=chapter_index,
                    active_chapter_title=chunk["chapter_title"],
                    chapter_chunk_index=chapter_chunk_positions[chapter_index],
                    chapter_chunk_count=chapter_chunk_totals.get(chapter_index, 0),
                    selected_chapter_indexes=selected_chapter_indexes,
                ),
            )

            segments = [SegmentAnnotation.model_validate(segment) for segment in chunk["segments"]]
            prompt = build_gemini_tts_prompt(
                book_title=annotations.title,
                chapter_title=chunk["chapter_title"],
                voice_name=request.voice_name,
                language_code=request.language_code,
                style_instruction=request.style_instruction,
                segments=segments,
            )
            def on_retry(next_attempt: int, max_attempts: int, reason: str) -> None:
                write_audio_generation_status(
                    book_id,
                    build_audio_status_payload(
                        state="running",
                        step=f"Retrying chunk {position} of {len(chunks)}: attempt {next_attempt}/{max_attempts}",
                        chunk_index=position,
                        chunk_count=len(chunks),
                        active_chapter_index=chapter_index,
                        active_chapter_title=chunk["chapter_title"],
                        chapter_chunk_index=chapter_chunk_positions[chapter_index],
                        chapter_chunk_count=chapter_chunk_totals.get(chapter_index, 0),
                        selected_chapter_indexes=selected_chapter_indexes,
                        retry_reason=reason,
                    ),
                )

            audio_bytes, sample_rate, attempts_used = synthesize_chunk_audio_with_gemini(
                prompt,
                request.voice_name,
                request.language_code,
                on_retry=on_retry,
            )
            chapter_audio.setdefault(chunk["chapter_index"], []).append(audio_bytes)
            chapter_sample_rates.setdefault(chunk["chapter_index"], sample_rate)
            chapter_titles[chunk["chapter_index"]] = chunk["chapter_title"]

            chunk_path = build_chunk_audio_path(book_id, chunk["chapter_index"], chunk["chunk_index"], request.voice_name)
            write_wave_file(chunk_path, audio_bytes, rate=sample_rate)

            if chapter_chunk_positions[chapter_index] == chapter_chunk_totals.get(chapter_index, 0):
                write_audio_generation_status(
                    book_id,
                    build_audio_status_payload(
                        state="running",
                        step=f"Finalizing chapter {chapter_index + 1}",
                        chunk_index=position,
                        chunk_count=len(chunks),
                        active_chapter_index=chapter_index,
                        active_chapter_title=chunk["chapter_title"],
                        chapter_chunk_index=chapter_chunk_positions[chapter_index],
                        chapter_chunk_count=chapter_chunk_totals.get(chapter_index, 0),
                        selected_chapter_indexes=selected_chapter_indexes,
                    ),
                )
                chapter_output_path = materialize_chapter_audio_output(
                    book_id,
                    chapter_index,
                    chapter_titles[chapter_index],
                    request.voice_name,
                    request.speech_rate,
                    request.output_format,
                    request.audio_quality,
                    chapter_audio[chapter_index],
                    chapter_sample_rates.get(chapter_index, DEFAULT_AUDIO_SAMPLE_RATE),
                )
                chapter_output_paths.append(chapter_output_path)
                generated_chapters.append(
                    {
                        "chapter_index": chapter_index,
                        "chapter_title": chapter_titles[chapter_index],
                        "file_name": chapter_output_path.name,
                        "download_url": f"/files/audio/{book_id}/{chapter_output_path.name}",
                        "output_format": request.output_format,
                    }
                )
                chapter_audio.pop(chapter_index, None)
                chapter_sample_rates.pop(chapter_index, None)
                chapter_titles.pop(chapter_index, None)

        write_audio_generation_status(
            book_id,
            build_audio_status_payload(
                state="running",
                step="Merging audio chunks",
                chunk_index=len(chunks),
                chunk_count=len(chunks),
                selected_chapter_indexes=selected_chapter_indexes,
            ),
        )

        if len(chapter_output_paths) == 1:
            final_download_path = chapter_output_paths[0]
        else:
            write_audio_generation_status(
                book_id,
                build_audio_status_payload(
                    state="running",
                    step="Packaging chapter MP3 files",
                    chunk_index=len(chunks),
                    chunk_count=len(chunks),
                    selected_chapter_indexes=selected_chapter_indexes,
                ),
            )
            final_download_path = build_audio_zip_path(book_id, request.voice_name, request.speech_rate, request.output_format)
            with zipfile.ZipFile(final_download_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
                for chapter_output_path in chapter_output_paths:
                    archive.write(chapter_output_path, arcname=chapter_output_path.name)

        write_audio_generation_status(
            book_id,
            build_audio_status_payload(
                state="completed",
                step="Completed",
                chunk_index=len(chunks),
                chunk_count=len(chunks),
                download_url=f"/files/audio/{book_id}/{final_download_path.name}",
                file_name=final_download_path.name,
                selected_chapter_indexes=selected_chapter_indexes,
                coverage_check=coverage_audit,
            ),
        )
    except Exception as exc:
        detail = exc.detail if isinstance(exc, HTTPException) else str(exc)
        selected_chapter_indexes = request.chapter_indexes or []
        write_audio_generation_status(
            book_id,
            {
                "state": "error",
                "step": detail or "Audio generation failed",
                "chunk_index": 0,
                "chunk_count": 0,
                "download_url": None,
                "file_name": None,
                "speech_rate": request.speech_rate,
                "voice_name": request.voice_name,
                "language_code": request.language_code,
                "style_instruction": request.style_instruction,
                "output_format": request.output_format,
                "audio_quality": request.audio_quality,
                "selected_chapter_indexes": selected_chapter_indexes,
                "generated_chapters": generated_chapters,
                "updated_at": now_iso(),
            },
        )


def compute_draft_annotation_progress_percent(phase: str, completed_chapters: int, total_chapters: int) -> int:
    if phase == "Preparing chapters":
        return 5
    if phase == "Annotating chapters":
        if total_chapters <= 0:
            return 15
        return min(90, 15 + int((completed_chapters / total_chapters) * 70))
    if phase == "Writing DOCX":
        return 95
    if phase == "Completed":
        return 100
    return 0


def run_draft_annotation_generation(book_id: str, request: "DraftAnnotationRequest") -> None:
    try:
        book = get_book_or_404(book_id)
        extracted_path = Path(book["extracted_path"])
        extracted = read_json(extracted_path)

        def on_progress(phase: str, completed_chapters: int, total_chapters: int, current_chapter_title: str | None) -> None:
            step = phase
            chapter_index = completed_chapters
            if phase == "Annotating chapters" and total_chapters > 0:
                current_number = min(completed_chapters + 1, total_chapters) if completed_chapters < total_chapters else total_chapters
                step = f"Annotating chapter {current_number} of {total_chapters}"
            write_draft_annotation_status(
                book_id,
                {
                    "state": "running",
                    "phase": phase,
                    "step": step,
                    "chapter_index": chapter_index,
                    "chapter_count": total_chapters,
                    "current_chapter_title": current_chapter_title,
                    "provider": "gemini" if get_gemini_api_key() else "heuristic",
                    "expressive_mode": request.expressive_mode,
                    "vocalization_mode": request.vocalization_mode,
                    "progress_percent": compute_draft_annotation_progress_percent(phase, completed_chapters, total_chapters),
                    "download_url": None,
                    "updated_at": now_iso(),
                },
            )

        annotations, provider = generate_annotation_document(
            extracted,
            request.expressive_mode,
            request.vocalization_mode,
            on_progress,
        )
        extracted["generated_annotation"] = annotations.model_dump()
        extracted["generated_annotation_provider"] = provider
        extracted["generated_annotation_settings"] = {
            "expressive_mode": request.expressive_mode,
            "vocalization_mode": request.vocalization_mode,
        }
        write_json(extracted_path, extracted)
        draft_path = create_draft_annotation_docx(
            book,
            annotations,
            request.expressive_mode,
            request.vocalization_mode,
        )
        timestamp = now_iso()
        with get_connection() as connection:
            connection.execute(
                "UPDATE books SET draft_annotation_path = ?, status = ?, updated_at = ? WHERE id = ?",
                (str(draft_path), "draft_annotation_ready", timestamp, book_id),
            )
        write_draft_annotation_status(
            book_id,
            {
                "state": "completed",
                "phase": "Completed",
                "step": "Completed",
                "chapter_index": len(annotations.chapters),
                "chapter_count": len(annotations.chapters),
                "current_chapter_title": None,
                "provider": provider,
                "expressive_mode": request.expressive_mode,
                "vocalization_mode": request.vocalization_mode,
                "progress_percent": 100,
                "download_url": f"/api/books/{book_id}/draft-annotation/download",
                "updated_at": timestamp,
            },
        )
    except Exception as exc:
        detail = exc.detail if isinstance(exc, HTTPException) else str(exc)
        with get_connection() as connection:
            connection.execute(
                "UPDATE books SET status = ?, updated_at = ? WHERE id = ?",
                ("draft_annotation_error", now_iso(), book_id),
            )
        write_draft_annotation_status(
            book_id,
            {
                "state": "error",
                "phase": "Error",
                "step": detail or "Draft annotation failed",
                "chapter_index": 0,
                "chapter_count": 0,
                "current_chapter_title": None,
                "provider": None,
                "expressive_mode": request.expressive_mode,
                "vocalization_mode": request.vocalization_mode,
                "progress_percent": 0,
                "download_url": None,
                "updated_at": now_iso(),
            },
        )


@app.on_event("startup")
def startup() -> None:
    init_db()


ensure_storage()
app.mount("/files/annotated", StaticFiles(directory=ANNOTATED_DIR), name="annotated")
app.mount("/files/audio", StaticFiles(directory=AUDIO_DIR), name="audio")


@app.get("/", response_class=HTMLResponse)
def index() -> str:
    return INDEX_PATH.read_text(encoding="utf-8")


@app.get("/api/health")
def health() -> dict[str, str]:
    return {
        "status": "ok",
        "gemini_configured": "true" if get_gemini_api_key() else "false",
        "gemini_text_model": get_gemini_text_model(),
        "gemini_tts_model": get_gemini_tts_model(),
    }


@app.get("/api/generator-options")
def get_generator_options() -> dict[str, Any]:
    return {
        "default_settings": build_default_generator_settings(),
        "voices": SUPPORTED_GEMINI_VOICES,
        "languages": SUPPORTED_LANGUAGE_OPTIONS,
        "chunk_limit_presets": CHUNK_TARGET_PRESETS,
        "hard_limit_overflow_presets": CHUNK_HARD_OVERFLOW_PRESETS,
        "output_formats": ["mp3", "wav"],
        "audio_quality_presets": ["best", "standard"],
    }


@app.get("/api/books")
def list_books() -> list[dict[str, Any]]:
    with get_connection() as connection:
        rows = connection.execute("SELECT * FROM books ORDER BY created_at DESC").fetchall()
    return [serialize_book(row) for row in rows]


@app.get("/api/books/{book_id}")
def get_book(book_id: str) -> dict[str, Any]:
    return serialize_book(get_book_or_404(book_id))


@app.post("/api/upload")
def upload_book(file: UploadFile = File(...)) -> dict[str, Any]:
    ensure_storage()
    source_path, suffix = save_upload(file, ALLOWED_SOURCE_EXTENSIONS, UPLOADS_DIR)
    book_id = uuid.uuid4().hex
    title = sanitize_stem(file.filename or "book").replace("-", " ").title()
    source_type = suffix.lstrip(".")

    extracted_source = extract_pdf(source_path) if suffix == ".pdf" else extract_docx(source_path)
    extracted_payload = build_extracted_payload(
        book_id,
        title,
        source_type,
        extracted_source.get("document_title"),
        extracted_source["paragraphs"],
    )
    extracted_path = EXTRACTED_DIR / f"{book_id}.json"
    write_json(extracted_path, extracted_payload)

    timestamp = now_iso()
    with get_connection() as connection:
        connection.execute(
            """
            INSERT INTO books (
                id, title, source_type, source_filename, source_path, extracted_path,
                status, chapter_count, paragraph_count, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                book_id,
                title,
                source_type,
                file.filename or source_path.name,
                str(source_path),
                str(extracted_path),
                "uploaded",
                extracted_payload["chapter_count"],
                extracted_payload["paragraph_count"],
                timestamp,
                timestamp,
            ),
        )

    return serialize_book(get_book_or_404(book_id))


@app.post("/api/books/{book_id}/draft-annotation")
def draft_annotation(
    book_id: str,
    request: DraftAnnotationRequest = Body(default=DraftAnnotationRequest()),
    background_tasks: BackgroundTasks = None,
) -> dict[str, Any]:
    book = get_book_or_404(book_id)
    existing_status = read_draft_annotation_status(book_id)
    if existing_status and existing_status.get("state") in {"queued", "running"}:
        raise HTTPException(status_code=409, detail="Draft annotation generation is already running for this book.")
    if background_tasks is None:
        raise HTTPException(status_code=500, detail="Background task runner is unavailable.")

    timestamp = now_iso()
    with get_connection() as connection:
        connection.execute(
            "UPDATE books SET status = ?, updated_at = ? WHERE id = ?",
            ("draft_annotation_running", timestamp, book_id),
        )
    write_draft_annotation_status(
        book_id,
        {
            "state": "queued",
            "phase": "Queued",
            "step": "Queued",
            "chapter_index": 0,
            "chapter_count": 0,
            "current_chapter_title": None,
            "provider": None,
            "expressive_mode": request.expressive_mode,
            "vocalization_mode": request.vocalization_mode,
            "progress_percent": 0,
            "download_url": None,
            "updated_at": timestamp,
        },
    )
    background_tasks.add_task(run_draft_annotation_generation, book_id, request)
    return {
        "status": "queued",
        "step": "Queued",
        "expressive_mode": request.expressive_mode,
        "vocalization_mode": request.vocalization_mode,
    }


@app.get("/api/books/{book_id}/draft-annotation-status")
def get_draft_annotation_status(book_id: str) -> dict[str, Any]:
    get_book_or_404(book_id)
    status = read_draft_annotation_status(book_id)
    if status is None:
        return {
            "state": "idle",
            "phase": "Idle",
            "step": "Idle",
            "chapter_index": 0,
            "chapter_count": 0,
            "current_chapter_title": None,
            "provider": None,
            "expressive_mode": None,
            "vocalization_mode": None,
            "progress_percent": 0,
            "download_url": None,
        }
    return status


@app.get("/api/books/{book_id}/draft-annotation/download")
def download_draft_annotation(book_id: str) -> FileResponse:
    book = get_book_or_404(book_id)
    if not book["draft_annotation_path"]:
        raise HTTPException(status_code=404, detail="No draft annotation available.")
    path = Path(book["draft_annotation_path"])
    if not path.exists():
        raise HTTPException(status_code=404, detail="Draft annotation file is missing.")
    return FileResponse(path, filename=build_annotated_download_name(book))


@app.post("/api/books/{book_id}/approved-annotation")
def upload_approved_annotation(book_id: str, file: UploadFile = File(...)) -> dict[str, str]:
    book = get_book_or_404(book_id)
    saved_path, _ = save_upload(file, ALLOWED_ANNOTATION_EXTENSIONS, ANNOTATED_DIR)
    annotations = parse_annotation_docx(saved_path)
    extracted = read_json(Path(book["extracted_path"]))
    extracted["approved_annotation"] = annotations.model_dump()
    write_json(Path(book["extracted_path"]), extracted)
    with get_connection() as connection:
        connection.execute(
            "UPDATE books SET approved_annotation_path = ?, status = ?, updated_at = ? WHERE id = ?",
            (str(saved_path), "approved_annotation_uploaded", now_iso(), book_id),
        )
    return {"status": "ok"}


@app.post("/api/upload-approved-annotation")
def upload_approved_annotation_as_book(file: UploadFile = File(...)) -> dict[str, Any]:
    ensure_storage()
    saved_path, _ = save_upload(file, ALLOWED_ANNOTATION_EXTENSIONS, ANNOTATED_DIR)
    annotations = parse_annotation_docx(saved_path)

    book_id = uuid.uuid4().hex
    extracted_payload = build_extracted_payload_from_annotation_document(book_id, annotations, file.filename or saved_path.name)
    extracted_path = EXTRACTED_DIR / f"{book_id}.json"
    write_json(extracted_path, extracted_payload)

    timestamp = now_iso()
    with get_connection() as connection:
        connection.execute(
            """
            INSERT INTO books (
                id, title, source_type, source_filename, source_path, extracted_path,
                status, draft_annotation_path, approved_annotation_path, chapter_count, paragraph_count, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                book_id,
                extracted_payload["title"],
                "docx",
                file.filename or saved_path.name,
                str(saved_path),
                str(extracted_path),
                "approved_annotation_uploaded",
                None,
                str(saved_path),
                extracted_payload["chapter_count"],
                extracted_payload["paragraph_count"],
                timestamp,
                timestamp,
            ),
        )

    return serialize_book(get_book_or_404(book_id))


@app.post("/api/books/{book_id}/chunk-preview")
def preview_chunks(book_id: str, request: ChunkPreviewRequest) -> dict[str, Any]:
    book = get_book_or_404(book_id)
    extracted = read_json(Path(book["extracted_path"]))
    approved = extracted.get("approved_annotation")
    if not approved:
        raise HTTPException(status_code=400, detail="Upload an approved annotation DOCX before previewing chunks.")
    annotations = AnnotationDocument.model_validate(approved)
    prepared_chapters = prepare_chapters_for_generation(
        annotations.chapters,
        request.read_chapter_titles,
        request.pre_chapter_text_mode,
    )
    chunks = chunk_chapters(prepared_chapters, request.max_chars, request.hard_max_chars)
    return {"chunks": chunks, "chapters": build_chapter_summaries_from_chapters(prepared_chapters)}


@app.post("/api/books/{book_id}/tts-prompt-preview")
def preview_tts_prompt(book_id: str, request: TtsPromptPreviewRequest) -> dict[str, Any]:
    book = get_book_or_404(book_id)
    validate_language_code(request.language_code)
    extracted = read_json(Path(book["extracted_path"]))
    approved = extracted.get("approved_annotation")
    if not approved:
        raise HTTPException(status_code=400, detail="Upload an approved annotation DOCX before previewing a TTS prompt.")
    annotations = AnnotationDocument.model_validate(approved)
    prepared_chapters = prepare_chapters_for_generation(
        annotations.chapters,
        request.read_chapter_titles,
        request.pre_chapter_text_mode,
    )
    chunks = chunk_chapters(prepared_chapters, request.max_chars, request.hard_max_chars)
    selected_chunk = resolve_chunk(chunks, request.chapter_index, request.chunk_index)
    segments = [SegmentAnnotation.model_validate(segment) for segment in selected_chunk["segments"]]
    prompt = build_gemini_tts_prompt(
        book_title=annotations.title,
        chapter_title=selected_chunk["chapter_title"],
        voice_name=request.voice_name,
        language_code=request.language_code,
        style_instruction=request.style_instruction,
        segments=segments,
    )
    return {
        "voice_name": request.voice_name,
        "language_code": request.language_code,
        "style_instruction": request.style_instruction,
        "output_format": request.output_format,
        "audio_quality": request.audio_quality,
        "speech_rate": request.speech_rate,
        "chapter_title": selected_chunk["chapter_title"],
        "chunk_index": selected_chunk["chunk_index"],
        "char_count": selected_chunk["char_count"],
        "prompt": prompt,
    }


@app.post("/api/books/{book_id}/generate-chunk-audio")
def generate_chunk_audio(book_id: str, request: GenerateChunkAudioRequest) -> dict[str, Any]:
    book = get_book_or_404(book_id)
    validate_language_code(request.language_code)
    extracted = read_json(Path(book["extracted_path"]))
    approved = extracted.get("approved_annotation")
    if not approved:
        raise HTTPException(status_code=400, detail="Upload an approved annotation DOCX before generating audio.")
    annotations = AnnotationDocument.model_validate(approved)
    prepared_chapters = prepare_chapters_for_generation(
        annotations.chapters,
        request.read_chapter_titles,
        request.pre_chapter_text_mode,
    )
    chunks = chunk_chapters(prepared_chapters, request.max_chars, request.hard_max_chars)
    selected_chunk = resolve_chunk(chunks, request.chapter_index, request.chunk_index)
    segments = [SegmentAnnotation.model_validate(segment) for segment in selected_chunk["segments"]]
    prompt = build_gemini_tts_prompt(
        book_title=annotations.title,
        chapter_title=selected_chunk["chapter_title"],
        voice_name=request.voice_name,
        language_code=request.language_code,
        style_instruction=request.style_instruction,
        segments=segments,
    )
    audio_bytes, sample_rate, attempts_used = synthesize_chunk_audio_with_gemini(prompt, request.voice_name, request.language_code)
    output_path = build_chunk_audio_path(book_id, request.chapter_index, request.chunk_index, request.voice_name).with_suffix(f".{request.output_format}")
    raw_output_path = build_chunk_audio_path(book_id, request.chapter_index, request.chunk_index, request.voice_name)
    write_wave_file(raw_output_path, audio_bytes, rate=sample_rate)
    if request.output_format == "mp3":
        encode_mp3_with_speed(raw_output_path, output_path, request.speech_rate, request.audio_quality)
    else:
        output_path = raw_output_path
    return {
        "voice_name": request.voice_name,
        "language_code": request.language_code,
        "style_instruction": request.style_instruction,
        "output_format": request.output_format,
        "audio_quality": request.audio_quality,
        "speech_rate": request.speech_rate,
        "attempts_used": attempts_used,
        "chapter_title": selected_chunk["chapter_title"],
        "chunk_index": selected_chunk["chunk_index"],
        "download_url": f"/files/audio/{book_id}/{output_path.name}",
        "file_name": output_path.name,
    }


@app.post("/api/books/{book_id}/generate-book-audio")
def generate_book_audio(book_id: str, request: GenerateBookAudioRequest, background_tasks: BackgroundTasks) -> dict[str, Any]:
    book = get_book_or_404(book_id)
    validate_language_code(request.language_code)
    extracted = read_json(Path(book["extracted_path"]))
    approved = extracted.get("approved_annotation")
    if not approved:
        raise HTTPException(status_code=400, detail="Upload an approved annotation DOCX before generating audio.")
    existing_status = read_audio_generation_status(book_id)
    if existing_status and existing_status.get("state") == "running":
        raise HTTPException(status_code=409, detail="Audio generation is already running for this book.")

    prepared_chapters = prepare_chapters_for_generation(
        AnnotationDocument.model_validate(approved).chapters,
        request.read_chapter_titles,
        request.pre_chapter_text_mode,
    )
    selected_chapter_indexes = normalize_selected_chapter_indexes(len(prepared_chapters), request.chapter_indexes)
    extracted["generator_settings"] = {
        "voice_name": request.voice_name,
        "language_code": request.language_code,
        "style_instruction": request.style_instruction,
        "speech_rate": request.speech_rate,
        "max_chars": request.max_chars,
        "hard_max_chars": request.hard_max_chars,
        "read_chapter_titles": request.read_chapter_titles,
        "pre_chapter_text_mode": request.pre_chapter_text_mode,
        "output_format": request.output_format,
        "audio_quality": request.audio_quality,
        "chapter_indexes": selected_chapter_indexes,
    }
    write_json(Path(book["extracted_path"]), extracted)

    write_audio_generation_status(
        book_id,
        {
            "state": "queued",
            "step": "Queued",
            "chunk_index": 0,
            "chunk_count": 0,
            "download_url": None,
            "file_name": None,
            "speech_rate": request.speech_rate,
            "voice_name": request.voice_name,
            "language_code": request.language_code,
            "style_instruction": request.style_instruction,
            "hard_max_chars": request.hard_max_chars,
            "read_chapter_titles": request.read_chapter_titles,
            "pre_chapter_text_mode": request.pre_chapter_text_mode,
            "output_format": request.output_format,
            "audio_quality": request.audio_quality,
            "selected_chapter_indexes": selected_chapter_indexes,
            "updated_at": now_iso(),
        },
    )
    background_tasks.add_task(run_book_audio_generation, book_id, request)
    return {
        "status": "queued",
        "step": "Queued",
        "speech_rate": request.speech_rate,
        "voice_name": request.voice_name,
        "language_code": request.language_code,
        "style_instruction": request.style_instruction,
        "read_chapter_titles": request.read_chapter_titles,
        "pre_chapter_text_mode": request.pre_chapter_text_mode,
        "output_format": request.output_format,
        "audio_quality": request.audio_quality,
        "selected_chapter_indexes": selected_chapter_indexes,
    }


@app.get("/api/books/{book_id}/audio-status")
def get_audio_generation_status(book_id: str) -> dict[str, Any]:
    get_book_or_404(book_id)
    status = read_audio_generation_status(book_id)
    if status is None:
        return {
            "state": "idle",
            "step": "Idle",
            "chunk_index": 0,
            "chunk_count": 0,
            "download_url": None,
            "file_name": None,
        }
    return status


@app.get("/api/books/{book_id}/approved-annotation")
def download_approved_annotation(book_id: str) -> FileResponse:
    book = get_book_or_404(book_id)
    if not book["approved_annotation_path"]:
        raise HTTPException(status_code=404, detail="No approved annotation uploaded.")
    path = Path(book["approved_annotation_path"])
    if not path.exists():
        raise HTTPException(status_code=404, detail="Approved annotation file is missing.")
    return FileResponse(path, filename=build_annotated_download_name(book))
